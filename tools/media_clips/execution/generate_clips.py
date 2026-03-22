import os
import datetime
import subprocess # For AppleScript
from gnews import GNews
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from trafilatura import bare_extraction
from newspaper import Article
import requests
import re
from difflib import SequenceMatcher
from difflib import SequenceMatcher
from html import escape
from dateutil import parser as date_parser
import pytz
try:
    from googlenewsdecoder import new_decoderv1
except ImportError:
    from googlenewsdecoder.decoderv1 import new_decoderv1

# --- CONFIGURATION (Default values, override via arguments) ---
DEFAULT_QUERIES = [
    '"Test" AND "Search"'
]

TRUSTED_SOURCES = [

    'economist.com', 'nytimes.com', 'washingtonpost.com', 'wsj.com', 'usatoday.com', 
    'politico.com', 'thehill.com', 'axios.com', 'bloomberg.com', 'reuters.com', 
    'apnews.com', 'cnn.com', 'foxnews.com', 'cnbc.com', 'cbsnews.com', 
    'nbcnews.com', 'abcnews.com', 'npr.org', 'ft.com', 'financialtimes.com'
]

# Explicitly blocked sources (to fix matching errors)
BLOCKED_SOURCES = [
    'cnbc tv18', 'cnbctv18', 'hindu', 'times of india', 'indian express', 
    'deccan herald', 'prothom alo', 'economic times' # Safety additions
]

# Email Config
EMAIL_RECIPIENTS = ["recipient@example.com"] # Update this with actual recipients
EMAIL_SENDER = "francesco.lampertico@gmail.com"
EMAIL_SUBJECT = "India Media Clips - {date}"

# --- HELPER FUNCTIONS ---
def set_font(run, bold=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Calibri')

def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    # Explicitly add color and underline to assurances standard link look even if style is missing
    # ESCAPE TEXT for XML
    safe_text = escape(text)
    hyperlink = parse_xml(r'<w:hyperlink r:id="{}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr><w:t>{}</w:t></w:r></w:hyperlink>'.format(r_id, safe_text))
    paragraph._p.append(hyperlink)
    return hyperlink

def create_email_draft(attachment_path, subject, body, recipients, sender=None):
    print("Creating email draft...")
    recipient_str = ", ".join(recipients)
    
    # Escape quotes for AppleScript
    subject_safe = subject.replace('"', '\\"')
    body_safe = body.replace('"', '\\"').replace('\n', '\\n')
    
    # AppleScript to create message
    sender_prop = f'sender:"{sender}", ' if sender else ""
    
    script = f'''
    set recipientList to {{{", ".join(['"' + r + '"' for r in recipients])}}}
    set theSubject to "{subject_safe}"
    set theBody to "{body_safe}"
    set theFile to POSIX file "{attachment_path}"
    
    tell application "Mail"
        set newMessage to make new outgoing message with properties {{{sender_prop}subject:theSubject, content:theBody & return & return, visible:true}}
        tell newMessage
            make new attachment with properties {{file name:theFile}} at after the last paragraph
            
            repeat with r in recipientList
                make new to recipient at end of to recipients with properties {{address:r}}
            end repeat
        end tell
        activate
    end tell
    '''
    
    try:
        subprocess.run(['osascript', '-e', script], check=True)
        print("Email draft created successfully in Mail.app")
    except subprocess.CalledProcessError as e:
        print(f"Failed to create email draft: {e}")

import argparse

def main():
    parser = argparse.ArgumentParser(description='Generate Generalized Media Clips')
    parser.add_argument('--topic', type=str, help='The main topic/title for the clips (e.g., "India Media Clips")', required=True)
    parser.add_argument('--queries', type=str, help='Comma-separated list of search queries', required=True)
    parser.add_argument('--period', type=str, help='GNews period (e.g. 12h, 24h, 7d)', default=None)
    parser.add_argument('--suffix', type=str, help='Filename suffix (e.g. Partial)', default='')
    parser.add_argument('--since', type=str, help='Filter articles since this timestamp (YYYY-MM-DD HH:MM)', default=None)
    parser.add_argument('--target-date', type=str, help='Override the target date for the file (YYYY-MM-DD)', default=None)
    parser.add_argument('--output-dir', type=str, help='Specific output directory', default='output')
    parser.add_argument('--email-sender', type=str, help='Sender email address', default=None)
    parser.add_argument('--email-recipient', type=str, help='Recipient email address (comma-separated)', default=None)
    parser.add_argument('--no-email', action='store_true', help='Skip email draft creation')
    parser.add_argument('--all-sources', action='store_true', help='Include all sources, not just trusted mainstream media')
    parser.add_argument('--custom-sources', type=str, help='Comma-separated list of custom trusted domains', default=None)
    args = parser.parse_args()

    # Parse queries
    queries = [q.strip() for q in args.queries.split(',')] if args.queries else DEFAULT_QUERIES

    # Source filtering mode
    use_all_sources = args.all_sources
    active_trusted_sources = TRUSTED_SOURCES
    if args.custom_sources:
        active_trusted_sources = [d.strip() for d in args.custom_sources.split(',') if d.strip()]
        print(f"Using custom sources: {active_trusted_sources}")


    # Parse constraints
    since_date = None
    if args.since:
        try:
            since_date = date_parser.parse(args.since)
            if since_date.tzinfo is None:
                # Assume local if not specified, or just compare naive
                pass
            print(f"Filtering articles published after: {since_date}")
        except Exception as e:
            print(f"Error parsing --since: {e}")

    # Determine Target Date
    if args.target_date:
        try:
            target_date_obj = date_parser.parse(args.target_date).date()
            print(f"Using target date: {target_date_obj}")
        except Exception as e:
             print(f"Error parsing target date, using today: {e}")
             target_date_obj = datetime.date.today()
    else:
        target_date_obj = datetime.date.today()

    print("Starting Morning Clips Search...")
    
    # Determine Period
    if args.period:
        period = args.period
        print(f"Using custom period: {period}")
    else:
        # Check if today is Monday (0 = Monday, 6 = Sunday)
        today = datetime.date.today()
        if today.weekday() == 0:
            print("It's Monday! Extending search period to 72h (Weekend Coverage).")
            period = '72h'
        else:
            period = '24h'

    # Increase max_results to capture more candidate articles
    google_news = GNews(language='en', country='US', period=period, max_results=50)
    
    found_articles = []
    seen_urls = set()
    seen_titles = set()

    # 1. SEARCH
    for q in queries:
        print(f"Searching: {q}...")

        results = google_news.get_news(q)
        for item in results:
            # Filter Sources
            source_url = item.get('url', '')
            source_name = item.get('publisher', {}).get('title', 'Unknown Source').lower()
            title = item.get('title', '').strip()
            
            # 0. Check BLOCKED Sources
            if any(blocked in source_name or blocked in source_url.lower() for blocked in BLOCKED_SOURCES):
                continue

            # Simple domain check
            is_trusted = use_all_sources  # If all sources, start as True
            for domain in active_trusted_sources:
                # 1. Check if full domain in URL (rarely true for GNews redirect links)
                if domain in source_url:
                    is_trusted = True
                    break
                
                # 2. Check if domain stem in Publisher Name
                # 'reuters.com' -> 'reuters'
                domain_stem = domain.split('.')[0]
                # Special case for FT, NYT, WSJ stems
                if domain_stem == 'ft':
                    if 'financial times' in source_name or ' ft ' in f" {source_name} " or 'ft.com' in source_name:
                        is_trusted = True
                        break
                elif domain_stem == 'nytimes':
                    if 'new york times' in source_name or 'nytimes' in source_name:
                        is_trusted = True
                        break
                elif domain_stem == 'wsj':
                    if 'wall street journal' in source_name or 'wsj' in source_name:
                        is_trusted = True
                        break
                elif domain_stem == 'washingtonpost':
                    if 'washington post' in source_name:
                        is_trusted = True
                        break
                elif domain_stem in source_name.replace(' ', ''):
                    is_trusted = True
                    break
            
            # Deduplicate (URL + Title)
            # Normalize title for dedup
            norm_title = re.sub(r'\W+', '', title.lower())
            
            # Date check
            article_date_str = item.get('published date', '')
            pass_date_filter = True
            if since_date and article_date_str:
                try:
                    art_date = date_parser.parse(article_date_str)
                    # Handle tz awareness mismatch
                    if since_date.tzinfo and not art_date.tzinfo:
                         art_date = art_date.replace(tzinfo=since_date.tzinfo) # Assume same
                    elif not since_date.tzinfo and art_date.tzinfo:
                         art_date = art_date.replace(tzinfo=None)
                    
                    if art_date < since_date:
                        pass_date_filter = False
                except:
                    pass

            if is_trusted and pass_date_filter and source_url not in seen_urls and norm_title not in seen_titles:
                seen_urls.add(source_url)
                seen_titles.add(norm_title)
                found_articles.append(item)

    print(f"Found {len(found_articles)} potential articles. Processing content...")

    # 1.5 PROCESS & FILTER CONTENT
    final_articles = []
    accepted_bodies = []
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8'
    }

    for article in found_articles:
        url = article.get('url', '')
        title = article.get('title', '').strip()
        
        # 1. Resolve URL
        try:
             decoded_obj = new_decoderv1(url)
             if decoded_obj.get("status"):
                 final_url = decoded_obj["decoded_url"]
             else:
                 final_url = url
        except Exception:
             final_url = url
        
        # 2. Extract content
        extracted_text = ""
        extracted_author = None
        
        try:
            response = requests.get(final_url, headers=headers, timeout=10)
            if response.status_code == 200:
                html_content = response.text
                
                # A. Extract Text (Trafilatura)
                extracted_data = bare_extraction(html_content)
                if extracted_data:
                    extracted_text = getattr(extracted_data, 'text', '')
                    extracted_author = getattr(extracted_data, 'author', None)
                
                # B. Extract Author (Newspaper3k)
                try:
                    article_np = Article(final_url)
                    article_np.set_html(html_content)
                    article_np.parse()
                    if article_np.authors:
                        extracted_author = ', '.join(article_np.authors)
                except Exception:
                    pass
        except Exception as e_req:
            print(f"Request error for {final_url}: {e_req}")
            pass

        # 3. Content Deduplication
        # Use first 1000 chars as signature
        if extracted_text and len(extracted_text) > 100:
            current_sig = extracted_text[:1000]
            is_dup_body = False
            for past_sig in accepted_bodies:
                # If > 80% similar to ANY previous article
                if SequenceMatcher(None, current_sig, past_sig).ratio() > 0.80:
                    is_dup_body = True
                    break
            
            if is_dup_body:
                print(f"Skipping duplicate content for: {title}")
                continue # Skip this article entirely
            
            # Not a dup, accept it
            accepted_bodies.append(current_sig)
        
        # Store processed data in article dict
        article['final_url'] = final_url
        article['extracted_text'] = extracted_text
        article['extracted_author'] = extracted_author
        
        final_articles.append(article)

    print(f"Final Count after Content Dedup: {len(final_articles)}")

    # 2. GENERATE DOC
    doc = Document()
    email_body_text = "" # Accumulate text for email
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    today_str = target_date_obj.strftime("%B %d, %Y")

    # A. TITLE PAGE
    # email_body_text += title_text # REMOVED: Custom greeting handling at end
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t1 = p_title.add_run(f"{args.topic}\n")
    set_font(run_t1, bold=True)

    run_t2 = p_title.add_run(today_str)
    set_font(run_t2, bold=True)

    doc.add_paragraph() # Spacer

    # B. INDEX (Summary List)
    for idx, article in enumerate(final_articles, 1):
        source = article.get('publisher', {}).get('title', 'Source')
        title = article.get('title', 'No Title')
        url = article.get('final_url', article.get('url', '')) # Use final URL
        date = article.get('published date', today_str)

        # Clean inputs
        if source.lower().endswith('.com'):
            source = source[:-4]
        
        # Title Case Source
        source = source.title()
        
        # Parse Date
        date_cleaned = date
        try:
             from dateutil import parser
             dt = parser.parse(date)
             date_cleaned = dt.strftime("%B %d, %Y")
        except:
             try:
                 parts = date.split(' ')
                 if len(parts) >= 4:
                     day = parts[1]
                     month_short = parts[2]
                     year = parts[3]
                     month_map = {"Jan": "January", "Feb": "February", "Mar": "March", "Apr": "April", "May": "May", "Jun": "June", 
                                  "Jul": "July", "Aug": "August", "Sep": "September", "Oct": "October", "Nov": "November", "Dec": "December"}
                     month_long = month_map.get(month_short, month_short)
                     date_cleaned = f"{month_long} {day}, {year}"
             except:
                 pass

        # Clean Title (Remove trailing "- Source")
        separator = f" - {source}"
        if separator.lower() in title.lower():
             idx_sep = title.lower().rfind(separator.lower())
             if idx_sep != -1:
                 title = title[:idx_sep]
        
        # Add to Email Body
        index_line = f"{idx}. {source}: {title} – {date_cleaned}\n{url}\n"
        email_body_text += index_line + "\n"

        p = doc.add_paragraph()
        run_idx = p.add_run(f"{idx}. {source}: ")
        set_font(run_idx, bold=True)
        add_hyperlink(p, url, title, bold=False)
        run_date = p.add_run(f" – {date_cleaned}")
        set_font(run_date, bold=False)

    doc.add_page_break()
    email_body_text += "\n" + "="*20 + "\n\n"

    # C. CLIPS (Details)
    for article in final_articles:
        source = article.get('publisher', {}).get('title', 'Source')
        if source.lower().endswith('.com'):
            source = source[:-4]
        
        # Title Case Source
        source = source.title()
            
        title = article.get('title', 'No Title')
        separator = f" - {source}"
        if separator.lower() in title.lower():
             idx_sep = title.lower().rfind(separator.lower())
             if idx_sep != -1:
                 title = title[:idx_sep]
                 
        url = article.get('final_url')
        date = article.get('published date', today_str)
        # Date Cleaning
        date_cleaned_clip = date
        try:
             from dateutil import parser
             dt = parser.parse(date)
             date_cleaned_clip = dt.strftime("%B %d, %Y")
        except:
             try:
                 parts = date.split(' ')
                 if len(parts) >= 4:
                     day = parts[1]
                     month_short = parts[2]
                     year = parts[3]
                     month_map = {"Jan": "January", "Feb": "February", "Mar": "March", "Apr": "April", "May": "May", "Jun": "June", 
                                  "Jul": "July", "Aug": "August", "Sep": "September", "Oct": "October", "Nov": "November", "Dec": "December"}
                     month_long = month_map.get(month_short, month_short)
                     date_cleaned_clip = f"{month_long} {day}, {year}"
             except:
                  if ',' in date:
                      date_cleaned_clip = date.split(',', 1)[1].strip()

        desc = article.get('description', '[No snippet available]')
        extracted_text = article.get('extracted_text', '')
        extracted_author = article.get('extracted_author', None)

        # 3. Header Block
        p_head = doc.add_paragraph()
        run_src = p_head.add_run(f"{source}\n")
        set_font(run_src, bold=True)
        add_hyperlink(p_head, url, f"{title}", bold=False)
        p_head.add_run("\n")
        
        author_str = "By Staff"
        if extracted_author:
            author_str = f"By {extracted_author}"
        
        run_auth = p_head.add_run(f"{author_str}\n")
        set_font(run_auth, bold=False)
        
        run_dt = p_head.add_run(f"{date_cleaned_clip}")
        set_font(run_dt, bold=False)
        
        # Email Body Header
        email_body_text += f"{source}\n{title}\n{url}\n{author_str}\n{date_cleaned_clip}\n\n"

        # 4. Body Text
        
        # Subtitle logic
        # Check against first paragraph to avoid duplication
        print_subtitle = True
        if extracted_text:
            paragraphs = extracted_text.split('\n')
            # Title removal heuristic
            start_idx = 0
            for i in range(min(3, len(paragraphs))):
                para_clean = paragraphs[i].strip()
                if not para_clean: continue
                t_clean = re.sub(r'[^\w\s]', '', title.lower())
                p_clean = re.sub(r'[^\w\s]', '', para_clean.lower())
                if t_clean in p_clean or p_clean in t_clean:
                    start_idx = i + 1
            
            # Check overlap with first actual paragraph
            # Find first non-empty paragraph after start_idx
            check_idx = start_idx
            while check_idx < len(paragraphs) and not paragraphs[check_idx].strip():
                check_idx += 1
            
            if check_idx < len(paragraphs):
                first_para = paragraphs[check_idx].strip()
                if desc and len(first_para) > 10 and SequenceMatcher(None, desc.lower(), first_para.lower()).ratio() > 0.8:
                    # Subtitle is redundant, don't print it.
                    print_subtitle = False
        
        if print_subtitle and desc and len(desc) > 10:
             # Title deduplication for subtitle
             similarity = SequenceMatcher(None, title.lower(), desc.lower()).ratio()
             if similarity < 0.8:
                 # Email
                 email_body_text += f"{desc}\n\n"
                 
                 p_sub = doc.add_paragraph(desc)
                 p_sub.paragraph_format.space_after = Pt(12)
                 for run in p_sub.runs:
                     set_font(run, bold=False)
                     run.italic = True

        if extracted_text:
            # Re-split (already done above, but index logic is local)
            paragraphs = extracted_text.split('\n')
            
            # Apply start_idx from above
            
            # C. Loop and Print
            first_printed = False
            seen_paragraphs = set() # Check ALL previous paragraphs
            
            for para in paragraphs[start_idx:]:
                para = para.strip()
                if not para:
                    continue
                
                # Intra-article Dedup: Check against ALL seen paragraphs
                # Use a normalized hash or just straight comparison
                # Normalize: remove whitespace, lower
                para_norm = re.sub(r'\s+', '', para.lower())
                
                if para_norm in seen_paragraphs:
                    continue
                
                # Also check similarity to *immediately previous* just in case of slight variance
                # (For speed, we only check last one for fuzzy, but all for exact)
                # Actually, GNews content often has exact dups.
                
                seen_paragraphs.add(para_norm)

                # Remove Dateline on the first paragraph
                if not first_printed:
                    para = re.sub(r'^[A-Za-z\s\(\)\.,]{3,50}\s?[-—–]\s?', '', para)
                    first_printed = True
                
                # CLEANUP: Remove "Photo Credit" and "First Published"
                if 'photo credit' in para.lower(): continue
                if 'first published:' in para.lower(): continue
                
                # Email
                email_body_text += f"{para}\n\n"

                p_body = doc.add_paragraph(para)
                p_body.paragraph_format.space_after = Pt(12) 
                for run in p_body.runs:
                    set_font(run, bold=False)
        else:
            # Fallback to description
            p_body = doc.add_paragraph("[PASTE FULL TEXT HERE]")
            for run in p_body.runs:
                set_font(run, bold=False)
            
        email_body_text += "\n" + "-"*10 + "\n\n"


    # Save
    base_output_dir = args.output_dir
    if not os.path.exists(base_output_dir):
        os.makedirs(base_output_dir)
        
    topic_dir = os.path.join(base_output_dir, args.topic)
    if not os.path.exists(topic_dir):
        os.makedirs(topic_dir)
    
    suffix_str = f"_{args.suffix}" if args.suffix else ""
    # Safe filename from topic
    safe_topic = re.sub(r'[^a-zA-Z0-9_\-]', '_', args.topic)
    filename = os.path.join(topic_dir, f"{safe_topic}_{target_date_obj.strftime('%B_%d')}{suffix_str}.docx")
    doc.save(filename)
    print(f"Saved to {filename}")
    
    # 3. Create Email Draft (only if email args provided and not skipped)
    if not args.no_email and args.email_sender and args.email_recipient:
        today_formatted = target_date_obj.strftime("%B %d")

        # Build email body with clips included
        email_body_final = f"Good morning,\n\nPlease find attached the {args.topic} for {today_formatted}.\n\n"
        email_body_final += email_body_text
        email_body_final += "\nBest regards"

        recipients_list = [r.strip() for r in args.email_recipient.split(',')]
        create_email_draft(
            attachment_path=filename,
            subject=f"{args.topic} - {today_str}",
            body=email_body_final,
            recipients=recipients_list,
            sender=args.email_sender
        )
    else:
        print("Email draft skipped (no email options provided or --no-email set).")


if __name__ == "__main__":
    main()