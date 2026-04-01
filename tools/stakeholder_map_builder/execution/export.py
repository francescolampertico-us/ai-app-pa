"""
Stakeholder Map — Exporter
============================
Produces two output files:
  - stakeholder_map.xlsx  (Sheet 1: Actors, Sheet 2: Relationships)
  - stakeholder_map.docx  (narrative summary organized by stance)
"""

import re
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Excel constants ───────────────────────────────────────────────────────────

ACTOR_COLUMNS = [
    ("Name",           22),
    ("Organization",   24),
    ("Type",           14),
    ("Stance",         14),
    ("Influence Tier", 14),
    ("Evidence",       50),
    ("Issue Areas",    40),
    ("LDA Amount ($)", 16),
    ("Source",         12),
    ("Notes",          30),
]

REL_COLUMNS = [
    ("From",              28),
    ("To",                28),
    ("Relationship Type", 20),
    ("Label",             36),
]

HEADER_FILL   = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_FONT   = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
BODY_FONT     = Font(name="Calibri", size=10)
ALT_ROW_FILL  = PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid")
THIN_BORDER   = Border(
    left=Side(style="thin", color="D0D0D0"),
    right=Side(style="thin", color="D0D0D0"),
    top=Side(style="thin", color="D0D0D0"),
    bottom=Side(style="thin", color="D0D0D0"),
)

# Stance fill colors (light tints for readability)
STANCE_FILLS = {
    "proponent": PatternFill(start_color="D5F5E3", end_color="D5F5E3", fill_type="solid"),
    "opponent":  PatternFill(start_color="FADBD8", end_color="FADBD8", fill_type="solid"),
    "neutral":   PatternFill(start_color="EAECEE", end_color="EAECEE", fill_type="solid"),
    "unknown":   PatternFill(start_color="D6EAF8", end_color="D6EAF8", fill_type="solid"),
}


def _write_header_row(ws, columns: list[tuple[str, int]]):
    for col_idx, (col_name, col_width) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER
        ws.column_dimensions[get_column_letter(col_idx)].width = col_width


def export_xlsx(result: dict, output_path: str):
    """Export stakeholder map to a formatted Excel file with two sheets."""
    wb = Workbook()

    # ── Sheet 1: Actors ───────────────────────────────────────────────────────
    ws_actors = wb.active
    ws_actors.title = "Actors"
    ws_actors.freeze_panes = "A2"

    _write_header_row(ws_actors, ACTOR_COLUMNS)

    actors = result.get("actors", [])
    id_to_name = {a["id"]: a["name"] for a in actors}

    for row_idx, actor in enumerate(actors, 2):
        stance = actor.get("stance", "unknown")
        lda = actor.get("lda_amount")
        try:
            lda_val = f"${float(lda):,.0f}" if lda is not None else ""
        except (ValueError, TypeError):
            lda_val = ""

        issue_areas = actor.get("issue_areas", [])
        issue_str = "; ".join(str(i)[:60] for i in issue_areas[:3])

        row_data = [
            actor.get("name", ""),
            actor.get("organization", ""),
            actor.get("stakeholder_type", "").title(),
            stance.title(),
            actor.get("influence_tier", "").title(),
            actor.get("evidence", ""),
            issue_str,
            lda_val,
            actor.get("source", ""),
            actor.get("notes", ""),
        ]

        for col_idx, value in enumerate(row_data, 1):
            cell = ws_actors.cell(row=row_idx, column=col_idx, value=str(value) if value else "")
            cell.font = BODY_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(vertical="top", wrap_text=True)

            # Color-code the Stance column (col 4)
            if col_idx == 4:
                stance_fill = STANCE_FILLS.get(stance)
                if stance_fill:
                    cell.fill = stance_fill
            elif row_idx % 2 == 0:
                cell.fill = ALT_ROW_FILL

        ws_actors.row_dimensions[row_idx].height = 28

    ws_actors.auto_filter.ref = (
        f"A1:{get_column_letter(len(ACTOR_COLUMNS))}{len(actors) + 1}"
    )

    # ── Sheet 2: Relationships ────────────────────────────────────────────────
    relationships = result.get("relationships", [])
    if relationships:
        ws_rels = wb.create_sheet("Relationships")
        ws_rels.freeze_panes = "A2"
        _write_header_row(ws_rels, REL_COLUMNS)

        for row_idx, rel in enumerate(relationships, 2):
            rel_data = [
                id_to_name.get(rel.get("from_id", ""), rel.get("from_id", "")),
                id_to_name.get(rel.get("to_id", ""), rel.get("to_id", "")),
                rel.get("type", "").replace("_", " ").title(),
                rel.get("label", ""),
            ]
            for col_idx, value in enumerate(rel_data, 1):
                cell = ws_rels.cell(row=row_idx, column=col_idx, value=str(value))
                cell.font = BODY_FONT
                cell.border = THIN_BORDER
                cell.alignment = Alignment(vertical="top")
                if row_idx % 2 == 0:
                    cell.fill = ALT_ROW_FILL
            ws_rels.row_dimensions[row_idx].height = 20

        ws_rels.auto_filter.ref = (
            f"A1:{get_column_letter(len(REL_COLUMNS))}{len(relationships) + 1}"
        )

    # ── Meta sheet ────────────────────────────────────────────────────────────
    ws_meta = wb.create_sheet("Info")
    meta_rows = [
        ("Field", "Value"),
        ("Policy Issue", result.get("policy_issue", "")),
        ("Generated At", result.get("generated_at", "")),
        ("Total Actors", str(len(actors))),
        ("Proponents", str(sum(1 for a in actors if a.get("stance") == "proponent"))),
        ("Opponents",  str(sum(1 for a in actors if a.get("stance") == "opponent"))),
        ("Relationships", str(len(relationships))),
    ]
    for row_idx, (k, v) in enumerate(meta_rows, 1):
        ws_meta.cell(row=row_idx, column=1, value=k).font = Font(name="Calibri", bold=(row_idx == 1))
        ws_meta.cell(row=row_idx, column=2, value=v).font = Font(name="Calibri")
    ws_meta.column_dimensions["A"].width = 20
    ws_meta.column_dimensions["B"].width = 50

    wb.save(output_path)


# ── DOCX export ───────────────────────────────────────────────────────────────

FONT_NAME     = "Calibri"
BODY_SIZE     = 11
HEADING_SIZE  = 14
SUB_SIZE      = 12


def _set_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def _add_heading(doc, text, level=1):
    size = HEADING_SIZE if level == 1 else SUB_SIZE
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    p.space_after = Pt(4)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run)
    return p


def export_docx(result: dict, output_path: str):
    """Export stakeholder map narrative to a DOCX file."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_SIZE)

    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("STAKEHOLDER MAP")
    _set_font(run, size=16, bold=True)

    issue_p = doc.add_paragraph()
    issue_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = issue_p.add_run(result.get("policy_issue", ""))
    _set_font(run, size=13, bold=True)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(8)
    run = meta_p.add_run(f"Generated: {result.get('generated_at', '')}  |  PA AI Toolkit")
    _set_font(run, size=9, italic=True)

    # ── Issue Overview ────────────────────────────────────────────────────────
    if result.get("issue_summary"):
        _add_heading(doc, "Issue Overview")
        _add_body(doc, result["issue_summary"])

    actors = result.get("actors", [])
    id_to_name = {a["id"]: a["name"] for a in actors}

    proponents = [a for a in actors if a.get("stance") == "proponent"]
    opponents  = [a for a in actors if a.get("stance") == "opponent"]
    neutral    = [a for a in actors if a.get("stance") in ("neutral", "unknown")]

    def _actor_bullet(doc, actor):
        name  = actor.get("name", "")
        atype = actor.get("stakeholder_type", "other").title()
        org   = actor.get("organization", "")
        tier  = actor.get("influence_tier", "")
        evid  = actor.get("evidence", "")

        meta = atype
        if org and org != name:
            meta += f", {org}"
        tier_tag = f" [{tier} influence]" if tier else ""

        text = f"**{name}** ({meta}){tier_tag}"
        if evid:
            text += f" — {evid}"
        _add_bullet(doc, text)

    # ── Proponents ────────────────────────────────────────────────────────────
    if proponents:
        _add_heading(doc, f"Proponents ({len(proponents)})")
        if result.get("proponent_summary"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(result["proponent_summary"])
            _set_font(run, italic=True)
        for actor in proponents:
            _actor_bullet(doc, actor)

    # ── Opponents ─────────────────────────────────────────────────────────────
    if opponents:
        _add_heading(doc, f"Opponents ({len(opponents)})")
        if result.get("opponent_summary"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(result["opponent_summary"])
            _set_font(run, italic=True)
        for actor in opponents:
            _actor_bullet(doc, actor)

    # ── Neutral / Unknown ─────────────────────────────────────────────────────
    if neutral:
        _add_heading(doc, f"Neutral / Unknown ({len(neutral)})")
        for actor in neutral:
            name  = actor.get("name", "")
            notes = actor.get("notes", "")
            text  = f"**{name}**"
            if notes:
                text += f" — {notes}"
            _add_bullet(doc, text)

    # ── Key Relationships ─────────────────────────────────────────────────────
    relationships = result.get("relationships", [])
    if relationships:
        _add_heading(doc, "Key Relationships")
        for rel in relationships[:15]:
            from_name = id_to_name.get(rel["from_id"], rel["from_id"])
            to_name   = id_to_name.get(rel["to_id"],   rel["to_id"])
            label     = rel.get("label", rel.get("type", "").replace("_", " "))
            _add_bullet(doc, f"**{from_name}** {label} **{to_name}**")

    # ── Key Coalitions ────────────────────────────────────────────────────────
    coalitions = result.get("key_coalitions", [])
    if coalitions:
        _add_heading(doc, "Key Coalitions")
        for c in coalitions:
            _add_bullet(doc, c)

    # ── Strategic Notes ───────────────────────────────────────────────────────
    if result.get("strategic_notes"):
        _add_heading(doc, "Strategic Notes")
        _add_body(doc, result["strategic_notes"])

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "Stance classifications are LLM-inferred from public data — verify before strategic use."
    )
    _set_font(run, size=8, italic=True)

    doc.save(output_path)
