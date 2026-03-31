"""
DOCX Exporter for Stakeholder Briefing
========================================
Produces a professional one-pager Word document.
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


FONT_NAME = "Calibri"
BODY_SIZE = 11
HEADING_SIZE = 14
SUBHEADING_SIZE = 12


def _set_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def _add_heading(doc, text, level=1):
    size = HEADING_SIZE if level == 1 else SUBHEADING_SIZE
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    p.space_after = Pt(4)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    # Handle **bold** markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run)
    return p


def export_docx(result: dict, output_path: str):
    """Export a stakeholder briefing result to DOCX."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_SIZE)

    # Narrow margins for one-pager
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    h = result["header"]

    # --- Title ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("STAKEHOLDER BRIEFING")
    _set_font(run, size=16, bold=True)

    # --- Header info ---
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(h["stakeholder_name"])
    _set_font(run, size=13, bold=True)

    if h.get("organization"):
        org_p = doc.add_paragraph()
        org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = org_p.add_run(h["organization"])
        _set_font(run, size=11)

    # Meta line
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(8)
    run = meta_p.add_run(
        f"Meeting Purpose: {h['meeting_purpose']}  |  "
        f"Prepared: {h['date_prepared']}  |  By: {h['prepared_by']}"
    )
    _set_font(run, size=9, italic=True)

    # --- Profile ---
    profile = result.get("profile", {})
    if profile:
        _add_heading(doc, "Profile", level=1)
        if profile.get("summary"):
            _add_body(doc, profile["summary"])
        if profile.get("current_role"):
            _add_body(doc, f"**Current Role:** {profile['current_role']}")
        if profile.get("key_areas"):
            _add_body(doc, f"**Key Policy Areas:** {', '.join(profile['key_areas'])}")
        if profile.get("notable_positions"):
            _add_body(doc, f"**Notable Positions:** {profile['notable_positions']}")

    # --- Policy Positions ---
    positions = result.get("policy_positions", [])
    if positions:
        _add_heading(doc, "Policy Positions", level=1)
        for p in positions:
            text = p["position"]
            if p.get("evidence"):
                text += f" — {p['evidence']}"
            _add_bullet(doc, text)

    # --- Disclosure Data ---
    disclosures = result.get("disclosures", {})
    has_lda_entity = bool(disclosures.get("lda_entity"))
    has_lda_topic = bool(disclosures.get("lda_topic"))
    fara = disclosures.get("fara", {})
    has_fara = bool(fara.get("registrants") if isinstance(fara, dict) else fara)
    irs = disclosures.get("irs990", {})
    has_irs = bool(irs.get("organizations") if isinstance(irs, dict) else irs)

    if has_lda_entity or has_lda_topic or has_fara or has_irs:
        _add_heading(doc, "Disclosure Data", level=1)

        if has_lda_entity:
            _add_heading(doc, "LDA Lobbying (Stakeholder Activity)", level=2)
            for r in disclosures["lda_entity"][:5]:
                amount = r.get("amount_reported", "")
                try:
                    amount_str = f" (${float(amount):,.0f})" if amount and amount != "N/A" else ""
                except (ValueError, TypeError):
                    amount_str = ""
                _add_bullet(
                    doc,
                    f"**{r.get('registrant_name', 'N/A')}** for "
                    f"{r.get('client_name', 'N/A')}{amount_str} — "
                    f"Filed {r.get('filing_year', r.get('dt_posted', 'N/A'))}"
                )

        if has_lda_topic:
            _add_heading(doc, "Lobbying Activity on Meeting Topic", level=2)
            for r in disclosures["lda_topic"][:8]:
                amount = r.get("amount_reported", "")
                try:
                    amount_str = f" (${float(amount):,.0f})" if amount and amount != "N/A" else ""
                except (ValueError, TypeError):
                    amount_str = ""
                _add_bullet(
                    doc,
                    f"**{r.get('client_name', 'N/A')}** via "
                    f"{r.get('registrant_name', 'N/A')}{amount_str} — "
                    f"{r.get('filing_year', '')} {r.get('filing_period', '')}"
                )

        if has_fara:
            _add_heading(doc, "FARA Foreign Agent", level=2)
            regs = fara.get("registrants", []) if isinstance(fara, dict) else []
            fps = fara.get("foreign_principals", []) if isinstance(fara, dict) else []
            for r in regs[:5]:
                _add_bullet(
                    doc,
                    f"**{r.get('registrant_name', 'N/A')}** "
                    f"(Reg #{r.get('registration_number', 'N/A')})"
                )
            for fp in fps[:5]:
                _add_bullet(
                    doc,
                    f"Foreign Principal: **{fp.get('foreign_principal_name', 'N/A')}** "
                    f"({fp.get('state_or_country', 'N/A')})"
                )

        if has_irs:
            _add_heading(doc, "IRS 990 Nonprofit", level=2)
            orgs = irs.get("organizations", []) if isinstance(irs, dict) else []
            filings = irs.get("filings", []) if isinstance(irs, dict) else []
            for o in orgs[:3]:
                _add_bullet(
                    doc,
                    f"**{o.get('organization_name', 'N/A')}** "
                    f"(EIN: {o.get('ein', 'N/A')})"
                )
            for f in filings[:3]:
                rev = f.get("total_revenue", "")
                rev_str = f"${rev:,.0f}" if rev and rev != "N/A" else "N/A"
                _add_bullet(
                    doc,
                    f"Year {f.get('tax_year', 'N/A')}: Revenue {rev_str}"
                )

    # --- Recent News ---
    news = result.get("news", [])
    if news:
        _add_heading(doc, "Recent News", level=1)
        for n in news:
            _add_bullet(doc, f"{n['title']} — {n['source']} ({n.get('date', '')})")

    # --- Talking Points ---
    talking_points = result.get("talking_points", [])
    if talking_points:
        _add_heading(doc, "Suggested Talking Points", level=1)
        for i, tp in enumerate(talking_points, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{i}. {tp['point']}")
            _set_font(run, bold=True)
            if tp.get("rationale"):
                run = p.add_run(f"\n   {tp['rationale']}")
                _set_font(run, italic=True)

    # --- Key Questions ---
    questions = result.get("key_questions", [])
    if questions:
        _add_heading(doc, "Key Questions to Ask", level=1)
        for q in questions:
            _add_bullet(doc, f"**{q['question']}** — {q.get('purpose', '')}")

    # --- Footer ---
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("CONFIDENTIAL — FOR INTERNAL USE ONLY")
    _set_font(run, size=8, italic=True)

    doc.save(output_path)
