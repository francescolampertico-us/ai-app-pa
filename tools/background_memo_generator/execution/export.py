"""
Background Memo DOCX Exporter
================================
Produces a DOCX matching the exact house style observed in the example memos:

Structure:
  DATE: <date>          ← Normal style, plain
  SUBJECT: <subject>    ← Normal style, plain
  [blank line]
  Overview              ← Normal style, plain (section label)
  <overview paragraph>  ← Normal style, plain
  [blank line]
  Fast Facts            ← Normal style, plain (section label)
  • Bold sentence.      ← List Paragraph style, bold text
  • Bold sentence.
  [blank line]
  <Section Heading>     ← Normal style, plain
  [Sub-heading if any]  ← Normal style, bold
  <paragraph>           ← Normal style, plain
  ...
  Relevant Links        ← Normal style, plain
  • Label — URL         ← List Paragraph style, plain
"""

from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_normal(doc: Document, text: str, bold: bool = False) -> None:
    """Add a Normal-style paragraph."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = bold


def _add_list_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """Add a List Paragraph style bullet."""
    p = doc.add_paragraph(style="List Paragraph")
    run = p.add_run(text)
    run.bold = bold


def _add_blank(doc: Document) -> None:
    doc.add_paragraph(style="Normal")


def export_docx(result: dict, output_path: str, memo_date: str = "") -> None:
    """
    Write a background memo DOCX to output_path.

    Args:
        result:      Output from generator.generate_memo()
        output_path: Full path to write the .docx file
        memo_date:   Date string (default: today formatted as "Month DD, YYYY")
    """
    doc = Document()

    # Remove default top margin on first paragraph to match tight header style
    # (keep default Normal style otherwise)

    d = memo_date or date.today().strftime("%B %d, %Y")
    subject = result["subject"]

    # ── Header block ─────────────────────────────────────────────────────────
    _add_normal(doc, f"DATE:\t\t{d}")
    _add_normal(doc, f"SUBJECT:\t{subject} Background Memo")
    _add_blank(doc)

    # ── Overview ─────────────────────────────────────────────────────────────
    _add_normal(doc, "Overview")
    _add_normal(doc, result["overview"])
    _add_blank(doc)

    # ── Fast Facts ───────────────────────────────────────────────────────────
    _add_normal(doc, "Fast Facts")
    for fact in result["fast_facts"]:
        _add_list_paragraph(doc, fact, bold=True)
    _add_blank(doc)

    # ── Content sections ─────────────────────────────────────────────────────
    for section in result["sections"]:
        heading = section.get("heading", "")
        subsections = section.get("subsections", [])

        _add_normal(doc, heading)

        for sub in subsections:
            sub_heading = sub.get("heading")
            paragraphs = sub.get("paragraphs", [])

            if sub_heading:
                _add_normal(doc, sub_heading, bold=True)

            for para_text in paragraphs:
                _add_normal(doc, para_text)

        _add_blank(doc)

    # ── Relevant Links ───────────────────────────────────────────────────────
    _add_normal(doc, "Relevant Links")
    for link in result["links"]:
        label = link.get("label", "")
        url = link.get("url", "")
        _add_list_paragraph(doc, f"{label} — {url}")

    doc.save(output_path)
