"""
exporter.py — DOCX export with page footer support.

Mirrors the formatting observed in the 5 example memos:
- Lato Light font (with fallback)
- Justified alignment
- Bold + italic subsection headings for speakers
- Bold speaker names in body text
- Section headings at 14pt
- Page footer with confidentiality text on every page
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .config import (
    DOCX_FONT_NAME,
    DOCX_FONT_FALLBACK,
    DOCX_BODY_FONT_SIZE_PT,
    DOCX_SECTION_HEADING_SIZE_PT,
    DEFAULT_CONFIDENTIALITY_FOOTER,
    HEADING_HEARING_OVERVIEW,
    HEADING_QA,
)


def _set_run_font(run, font_name=DOCX_FONT_NAME, size_pt=DOCX_BODY_FONT_SIZE_PT):
    """Set font name and size on a run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # Set east-asian and other font families for compatibility
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
            f'w:cs="{font_name}" w:eastAsia="{font_name}"/>'
        )
        rPr.insert(0, rFonts)


def _add_paragraph(doc, text, font_name=DOCX_FONT_NAME,
                   size_pt=DOCX_BODY_FONT_SIZE_PT,
                   bold=False, italic=False,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_after=Pt(6)):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = Pt(0)

    run = para.add_run(text)
    _set_run_font(run, font_name, size_pt)
    run.bold = bold
    run.italic = italic

    return para


def _add_body_text_with_bold_speakers(doc, text, speaker_names=None):
    """Add body text with bold speaker names.

    In the example memos, speaker names at the start of sentences are bolded.
    """
    if not text:
        return

    # Split by double newlines for paragraph breaks
    paragraphs = text.split("\n\n")

    for para_text in paragraphs:
        if not para_text.strip():
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(0)

        # Find and bold speaker name references
        if speaker_names:
            _render_text_with_bold_names(para, para_text, speaker_names)
        else:
            run = para.add_run(para_text)
            _set_run_font(run)


def _render_text_with_bold_names(para, text, speaker_names):
    """Render paragraph text with specific names bolded."""
    # Build a pattern to match any speaker name
    if not speaker_names:
        run = para.add_run(text)
        _set_run_font(run)
        return

    # Escape names for regex and combine
    patterns = []
    for name in speaker_names:
        patterns.append(re.escape(name))
    combined = "|".join(sorted(patterns, key=len, reverse=True))
    regex = re.compile(f"({combined})")

    # Split text by speaker names
    parts = regex.split(text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part)
        _set_run_font(run)
        if regex.match(part):
            run.bold = True


def _collect_speaker_names(memo_output: dict) -> list:
    """Collect all speaker short-name references for bolding in body text."""
    names = set()

    for section in memo_output.get("sections", []):
        for sub in section.get("subsections", []):
            heading = sub.get("heading", "")
            # Extract last name from heading for short form
            parts = re.sub(r"\s*\([^)]*\)\s*$", "", heading).strip().split()
            if parts:
                last_name = parts[-1]
                # Build common short forms
                for prefix in ["Chairman", "Chairwoman", "Ranking Member",
                               "Sen.", "Rep.", "Commissioner", "Hon.",
                               "Mr.", "Ms.", "Dr."]:
                    candidate = f"{prefix} {last_name}"
                    names.add(candidate)
                # Also add full heading for subsection-heading matching
                names.add(heading)

    return list(names)


def export_docx(memo_output: dict, output_path: str) -> str:
    """Export the memo to a DOCX file with professional formatting.

    Args:
        memo_output: MemoOutput dict from the composer
        output_path: Path to write the DOCX file

    Returns:
        Path to the written file
    """
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # --- Page footer with confidentiality text ---
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = memo_output.get("footer", {}).get("text", DEFAULT_CONFIDENTIALITY_FOOTER)
    footer_run = footer_para.add_run(footer_text)
    _set_run_font(footer_run, DOCX_FONT_NAME, 9)
    footer_run.italic = True

    # Collect speaker names for bolding
    speaker_names = _collect_speaker_names(memo_output)

    # --- Metadata block (FROM / DATE / SUBJECT) ---
    mb = memo_output["metadata_block"]

    # Empty line before
    _add_paragraph(doc, "", size_pt=DOCX_BODY_FONT_SIZE_PT, space_after=Pt(0))

    # FROM line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run("FROM:\t\t\t")
    _set_run_font(run)
    run = para.add_run(mb["from"])
    _set_run_font(run)

    # DATE line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run("DATE:\t\t\t")
    _set_run_font(run)
    run = para.add_run(mb["date"])
    _set_run_font(run)

    # SUBJECT line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("SUBJECT:\t")
    _set_run_font(run)
    run = para.add_run(mb["subject"])
    _set_run_font(run)

    # --- Blank line ---
    _add_paragraph(doc, "", size_pt=DOCX_BODY_FONT_SIZE_PT, space_after=Pt(6))

    # --- Display title (centered) ---
    display_title = memo_output.get("display_title", "")
    if display_title:
        _add_paragraph(
            doc, display_title,
            size_pt=DOCX_BODY_FONT_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(12),
        )

    # --- Sections ---
    for section_data in memo_output.get("sections", []):
        heading = section_data.get("heading", "")
        body = section_data.get("body", "")
        subsections = section_data.get("subsections", [])

        # Section heading (14pt, Lato Light)
        _add_paragraph(
            doc, heading,
            size_pt=DOCX_SECTION_HEADING_SIZE_PT,
            space_after=Pt(6),
        )

        # Section body (for overview)
        if body:
            _add_body_text_with_bold_speakers(doc, body, speaker_names)

        # Subsections
        for sub in subsections:
            sub_heading = sub.get("heading", "")
            sub_body = sub.get("body", "")
            speaker_type = sub.get("speaker_type", "")

            # Subsection heading — bold + italic for leadership/members
            is_speaker = speaker_type in ("leadership_member", "qa_member")
            _add_paragraph(
                doc, sub_heading,
                bold=True,
                italic=is_speaker,
                space_after=Pt(4),
            )

            # Subsection body with bold speaker names
            _add_body_text_with_bold_speakers(doc, sub_body, speaker_names)

    # --- Confidentiality footer in content (inline, once) ---
    _add_paragraph(doc, "", space_after=Pt(12))
    _add_paragraph(
        doc, footer_text,
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(0),
    )

    # Save
    doc.save(output_path)
    return output_path
