"""
DOCX Exporter for Messaging Matrix
====================================
Produces a professional Word document with the Message House and all variants.
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from formats import VARIANT_LABELS


FONT_NAME = "Calibri"
BODY_SIZE = 11
HEADING_SIZE = 14
SUBHEADING_SIZE = 12


def _set_font(run, size=BODY_SIZE, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold


def _add_heading(doc, text, level=1):
    """Add a styled heading."""
    size = HEADING_SIZE if level == 1 else SUBHEADING_SIZE
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    p.space_after = Pt(6)
    return p


def _add_body(doc, text):
    """Add body text, handling basic markdown bold/italic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    # Split on **bold** markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run)
    return p


def _add_bullet(doc, text):
    """Add a bulleted paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    # Handle bold markers in bullets
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_font(run)
    return p


def _render_markdown_block(doc, content: str):
    """Parse a markdown block and add to document with basic formatting."""
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:], level=2)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=1)
        elif stripped.startswith("# "):
            _add_heading(doc, stripped[2:], level=1)
        # Bullets
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_bullet(doc, stripped[2:])
        # Numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            _add_bullet(doc, text)
        # Horizontal rule
        elif stripped.startswith("---"):
            doc.add_paragraph("").paragraph_format.space_after = Pt(12)
        # Regular text
        else:
            _add_body(doc, stripped)


def export_docx(result: dict, output_path: str):
    """
    Export a messaging matrix result to DOCX.

    Args:
        result: Output from generate_matrix() with message_house and variants.
        output_path: Path for the .docx file.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_SIZE)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    house = result["message_house"]

    # --- Title ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("MESSAGING MATRIX")
    _set_font(run, size=16, bold=True)

    doc.add_paragraph("")

    # --- Message Map ---
    _add_heading(doc, "Message Map", level=1)

    overarching = house.get("overarching_message", house.get("core_message", ""))
    key_messages = house.get("key_messages", house.get("pillars", []))

    if house.get("target_audiences"):
        p = doc.add_paragraph()
        run = p.add_run("Target Audiences: ")
        _set_font(run, bold=True)
        run = p.add_run(", ".join(house["target_audiences"]))
        _set_font(run)

    p = doc.add_paragraph()
    run = p.add_run("Overarching Message: ")
    _set_font(run, bold=True)
    run = p.add_run(overarching)
    _set_font(run)

    doc.add_paragraph("")

    # Build message map table
    if key_messages:
        num_cols = len(key_messages) + 1  # label column + one per key message
        max_facts = max(
            len(km.get("supporting_facts", km.get("proof_points", [])))
            for km in key_messages
        )
        num_rows = 1 + 1 + max_facts  # header + key message titles + fact rows

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        run = header_cells[0].paragraphs[0].add_run("")
        for ci, km in enumerate(key_messages, 1):
            run = header_cells[ci].paragraphs[0].add_run(f"Key Message {ci}")
            _set_font(run, bold=True)

        # Key message titles row
        title_cells = table.rows[1].cells
        run = title_cells[0].paragraphs[0].add_run("Key Message")
        _set_font(run, bold=True)
        for ci, km in enumerate(key_messages, 1):
            title = km.get("title", km.get("name", ""))
            run = title_cells[ci].paragraphs[0].add_run(title)
            _set_font(run)

        # Supporting facts rows
        for fi in range(max_facts):
            fact_cells = table.rows[2 + fi].cells
            run = fact_cells[0].paragraphs[0].add_run(f"Supporting Fact {fi + 1}")
            _set_font(run, bold=True)
            for ci, km in enumerate(key_messages, 1):
                facts = km.get("supporting_facts", km.get("proof_points", []))
                text = facts[fi] if fi < len(facts) else ""
                run = fact_cells[ci].paragraphs[0].add_run(text)
                _set_font(run)

    doc.add_paragraph("")

    if house.get("key_terms"):
        p = doc.add_paragraph()
        run = p.add_run("Key Terms: ")
        _set_font(run, bold=True)
        run = p.add_run(", ".join(house["key_terms"]))
        _set_font(run)

    # --- Variants ---
    for vid, content in result.get("variants", {}).items():
        doc.add_page_break()
        label = VARIANT_LABELS.get(vid, vid)
        _add_heading(doc, label, level=1)
        _render_markdown_block(doc, content)

    # --- Footer ---
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("CONFIDENTIAL — FOR INTERNAL REVIEW ONLY")
    _set_font(run, size=8, bold=False)
    run.font.italic = True

    doc.save(output_path)
