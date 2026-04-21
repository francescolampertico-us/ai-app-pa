from __future__ import annotations

import csv
import importlib.util
import json
import re
import subprocess
import sys
import threading
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, Request
from starlette.datastructures import UploadFile

from .jobs import create_job, set_job_artifacts, update_job

BMG_SCHEMA_PATH = Path(__file__).resolve().parents[3] / "toolkit" / "tools" / "background_memo" / "execution" / "schema.py"

router = APIRouter()

ROOT = Path(__file__).resolve().parents[3]
FASTAPI_ROOT = ROOT / "fastapi-backend"
TOOLKIT_ROOT = ROOT / "toolkit"
TOOLS_ROOT = TOOLKIT_ROOT / "tools"
JOBS_OUTPUT_ROOT = FASTAPI_ROOT / "jobs_output"
JOBS_OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
MEDIA_CLIPS_GENERATOR_PATH = TOOLS_ROOT / "media_clips" / "execution" / "generate_clips.py"

STYLE_GUIDES_DIR = TOOLS_ROOT / "messaging_matrix" / "style_samples" / "style_guides"
LEGISLATIVE_CACHE_DIR = TOOLS_ROOT / "legislative_tracker" / "execution" / ".cache"


def _ensure_sys_path(path: Path) -> None:
    path_str = str(path)
    if path_str not in sys.path:
        sys.path.insert(0, path_str)


def _load_module(name: str, path: Path):
    _ensure_sys_path(path.parent)
    spec = importlib.util.spec_from_file_location(name, path)
    if not spec or not spec.loader:
        raise RuntimeError(f"Could not load module at {path}")
    module = importlib.util.module_from_spec(spec)
    # Register before exec so forward-ref resolution (e.g. Pydantic model_rebuild)
    # can find the module via sys.modules[cls.__module__].
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def _job_output_dir(job_id: str) -> Path:
    out_dir = JOBS_OUTPUT_ROOT / job_id
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def _read_json(path: Path, default: Any = None) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8", errors="replace"))


def _read_text(path: Path, default: str = "") -> str:
    if not path.exists():
        return default
    return path.read_text(encoding="utf-8", errors="replace")


def _read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        return list(csv.DictReader(handle))


def _existing(paths: list[Path]) -> list[str]:
    seen: set[str] = set()
    existing: list[str] = []
    for path in paths:
        path_str = str(path)
        if path.exists() and path.is_file() and path_str not in seen:
            seen.add(path_str)
            existing.append(path_str)
    return existing


_media_clip_prepare_fn: Any | None = None


def _prepare_media_clip_body(text: str, title: str) -> list[str]:
    paragraphs: list[str] = []
    global _media_clip_prepare_fn

    if text:
        if _media_clip_prepare_fn is None:
            try:
                module = _load_module("_media_clips_generate_for_router", MEDIA_CLIPS_GENERATOR_PATH)
                prepare_fn = getattr(module, "_prepare_clip_paragraphs", None)
                _media_clip_prepare_fn = prepare_fn if callable(prepare_fn) else False
            except Exception:
                _media_clip_prepare_fn = False

        if _media_clip_prepare_fn:
            try:
                prepared = _media_clip_prepare_fn(text, title or "")
                if prepared:
                    paragraphs = [p.strip() for p in prepared if p and p.strip()]
            except Exception:
                paragraphs = []

    if paragraphs:
        return paragraphs
    return [p.strip() for p in text.split("\n") if p and p.strip()]


def _media_clips_date_stamp(raw_date: str | None) -> str:
    value = (raw_date or "").strip()
    if not value:
        return date.today().strftime("%b%d").lower()
    for fmt in ("%B %d, %Y", "%b %d, %Y", "%Y-%m-%d", "%m/%d/%Y"):
        try:
            return datetime.strptime(value, fmt).strftime("%b%d").lower()
        except ValueError:
            continue
    try:
        parsed = datetime.fromisoformat(value)
        return parsed.strftime("%b%d").lower()
    except ValueError:
        return date.today().strftime("%b%d").lower()


def _save_uploads(output_dir: Path, uploads: list[dict[str, Any]], field_name: str | None = None) -> list[Path]:
    saved: list[Path] = []
    for index, upload in enumerate(uploads):
        if field_name and upload["field"] != field_name:
            continue
        filename = Path(upload["filename"] or f"upload_{index}").name
        dest = output_dir / filename
        stem = dest.stem
        suffix = dest.suffix
        counter = 1
        while dest.exists():
            dest = output_dir / f"{stem}_{counter}{suffix}"
            counter += 1
        dest.write_bytes(upload["content"])
        saved.append(dest)
    return saved


def _extract_json_stdout(stdout: str) -> Any:
    text = (stdout or "").strip()
    if not text:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


CHANGE_AGENT_BASE_URL = "https://runpod-proxy-956966668285.us-central1.run.app/v1/"


def _subprocess_env(llm_model: str | None = None) -> dict[str, str] | None:
    """Return an env dict with OpenAI-compatible vars overridden for ChangeAgent, or None to inherit."""
    import os
    if llm_model != "ChangeAgent":
        return None
    env = os.environ.copy()
    env["OPENAI_BASE_URL"] = CHANGE_AGENT_BASE_URL
    env["OPENAI_API_KEY"] = os.environ.get("CHANGE_AGENT_API_KEY", "")
    env["CHANGE_AGENT_BASE_URL"] = CHANGE_AGENT_BASE_URL
    env["CHANGE_AGENT_API_KEY"] = os.environ.get("CHANGE_AGENT_API_KEY", "")
    env["LLM_MODEL_OVERRIDE"] = "ChangeAgent"
    return env


_llm_env_lock = threading.Lock()


from contextlib import contextmanager

@contextmanager
def _llm_env(llm_model: str | None):
    """Thread-safe context manager: temporarily override OPENAI_* env vars for ChangeAgent."""
    import os
    if llm_model != "ChangeAgent":
        yield
        return
    with _llm_env_lock:
        old_key = os.environ.get("OPENAI_API_KEY")
        old_base = os.environ.get("OPENAI_BASE_URL")
        old_model = os.environ.get("LLM_MODEL_OVERRIDE")
        try:
            os.environ["OPENAI_API_KEY"] = os.environ.get("CHANGE_AGENT_API_KEY", "")
            os.environ["OPENAI_BASE_URL"] = CHANGE_AGENT_BASE_URL
            os.environ["LLM_MODEL_OVERRIDE"] = "ChangeAgent"
            yield
        finally:
            if old_key is None:
                os.environ.pop("OPENAI_API_KEY", None)
            else:
                os.environ["OPENAI_API_KEY"] = old_key
            if old_base is None:
                os.environ.pop("OPENAI_BASE_URL", None)
            else:
                os.environ["OPENAI_BASE_URL"] = old_base
            if old_model is None:
                os.environ.pop("LLM_MODEL_OVERRIDE", None)
            else:
                os.environ["LLM_MODEL_OVERRIDE"] = old_model


def _run_command(
    job_id: str,
    cmd: list[str],
    cwd: Path,
    message: str,
    progress: int = 15,
    timeout: int = 600,
    llm_model: str | None = None,
) -> subprocess.CompletedProcess[str]:
    update_job(job_id, status="processing", progress=progress, message=message)
    result = subprocess.run(
        cmd,
        cwd=str(cwd),
        capture_output=True,
        text=True,
        timeout=timeout,
        env=_subprocess_env(llm_model),
    )
    if result.returncode != 0:
        detail = (result.stderr or result.stdout or "Tool execution failed").strip()
        raise RuntimeError(detail)
    return result


def _read_file_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    return path.read_text(encoding="utf-8", errors="replace").strip()


def _read_uploaded_context(upload_paths: list[Path]) -> str:
    parts: list[str] = []
    for path in upload_paths:
        try:
            text = _read_file_text(path)
        except Exception as exc:
            text = f"[Could not extract {path.name}: {exc}]"
        if text:
            parts.append(f"--- {path.name} ---\n{text[:8000]}")
    return "\n\n".join(parts)


def _fetch_youtube_transcript(url: str) -> str:
    match = re.search(r"(?:v=|youtu\.be/|/v/|/embed/)([a-zA-Z0-9_-]{11})", url)
    if not match:
        raise ValueError("Could not extract video ID from the YouTube URL.")

    from youtube_transcript_api import YouTubeTranscriptApi

    transcript = YouTubeTranscriptApi().fetch(match.group(1), languages=["en"])
    return "\n".join(snippet.text for snippet in transcript)


def _run_disclosure_tracker(
    job_id: str,
    output_dir: Path,
    *,
    entities: str,
    search_field: str,
    filing_years: str | None = None,
    all_years: bool = False,
    quarters: str = "Q1,Q2,Q3,Q4",
    from_date: str | None = None,
    to_date: str | None = None,
    sources: str = "lda,irs990",
    mode: str = "basic",
    max_results: str = "500",
    fuzzy_threshold: str = "85",
    dry_run: bool = False,
    llm_model: str | None = None,
) -> dict[str, Any]:
    run_path = TOOLS_ROOT / "influence_disclosure_tracker" / "execution" / "run.py"
    exec_dir = run_path.parent
    cmd = [
        sys.executable,
        str(run_path),
        "--entities",
        entities,
        "--search-field",
        search_field or "both",
        "--filing-periods",
        quarters or "Q1,Q2,Q3,Q4",
        "--sources",
        sources or "lda,irs990",
        "--out",
        str(output_dir),
        "--mode",
        mode or "basic",
        "--max-results",
        str(max_results or "500"),
        "--max-deep",
        "2",
        "--fuzzy-threshold",
        str(fuzzy_threshold or "85"),
        "--format",
        "csv,md",
    ]
    if not all_years and filing_years:
        cmd.extend(["--filing-years", filing_years])
    if from_date:
        cmd.extend(["--from", from_date])
    if to_date:
        cmd.extend(["--to", to_date])
    if dry_run:
        cmd.append("--dry-run")

    command_result = _run_command(job_id, cmd, exec_dir, "Querying disclosure databases...", progress=25, timeout=300, llm_model=llm_model)

    report_path = next(output_dir.rglob("report.md"), None)
    csv_paths = sorted(output_dir.rglob("*.csv"))
    result = {
        "report": _read_text(report_path) if report_path else "",
        "csv_data": {path.stem: _read_csv_rows(path) for path in csv_paths},
        "artifacts": _existing(([report_path] if report_path else []) + csv_paths),
        "stdout": command_result.stdout or "",
        "stderr": command_result.stderr or "",
    }
    return result


def _build_media_report(clips_data: list[dict[str, Any]], report_topic: str, report_date_str: str, output_dir: Path) -> tuple[Path, Path, Path]:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from html import escape

    def _set_font(run, bold: bool = False):
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), "Calibri")

    def _add_hyperlink(paragraph, url: str, text: str):
        part = paragraph.part
        rel_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = parse_xml(
            r'<w:hyperlink r:id="{}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            r'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            r'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0000FF"/>'
            r'<w:u w:val="single"/></w:rPr><w:t>{}</w:t></w:r></w:hyperlink>'.format(rel_id, escape(text))
        )
        paragraph._p.append(hyperlink)

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    email_body = ""
    html_body = '<div style="font-family:Calibri,sans-serif;font-size:14px;">\n'
    html_body += "<p>Good morning,</p>\n"
    html_body += f"<p>Please find attached the {report_topic} for {report_date_str}.</p>\n<hr>\n"

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = title_p.add_run(f"{report_topic}\n")
    _set_font(run1, bold=True)
    run2 = title_p.add_run(report_date_str)
    _set_font(run2, bold=True)
    doc.add_paragraph()

    for idx, clip in enumerate(clips_data, 1):
        index_p = doc.add_paragraph()
        prefix = index_p.add_run(f"{idx}. {clip.get('source', 'Unknown')}: ")
        _set_font(prefix, bold=True)
        if clip.get("url"):
            _add_hyperlink(index_p, clip["url"], clip.get("title", "Untitled"))
        else:
            title_run = index_p.add_run(clip.get("title", "Untitled"))
            _set_font(title_run)
        date_run = index_p.add_run(f" – {clip.get('date', '')}")
        _set_font(date_run)

        email_body += f"{idx}. {clip.get('source', 'Unknown')}: {clip.get('title', 'Untitled')} – {clip.get('date', '')}\n{clip.get('url', '')}\n\n"
        if clip.get("url"):
            html_body += f'<p><b>{idx}. {clip.get("source", "Unknown")}:</b> <a href="{clip["url"]}">{clip.get("title", "Untitled")}</a> &ndash; {clip.get("date", "")}</p>\n'
        else:
            html_body += f'<p><b>{idx}. {clip.get("source", "Unknown")}:</b> {clip.get("title", "Untitled")} &ndash; {clip.get("date", "")}</p>\n'

    doc.add_page_break()
    html_body += "<hr>\n"

    for clip in clips_data:
        header_p = doc.add_paragraph()
        src_run = header_p.add_run(f"{clip.get('source', 'Unknown')}\n")
        _set_font(src_run, bold=True)
        if clip.get("url"):
            _add_hyperlink(header_p, clip["url"], clip.get("title", "Untitled"))
        else:
            title_run = header_p.add_run(clip.get("title", "Untitled"))
            _set_font(title_run)
        header_p.add_run("\n")
        byline = clip.get("author") or "Staff"
        auth_run = header_p.add_run(f"By {byline}\n")
        _set_font(auth_run)
        date_run = header_p.add_run(clip.get("date", ""))
        _set_font(date_run)

        html_body += (
            f"<p><b>{clip.get('source', 'Unknown')}</b><br>\n"
            + (f'<a href="{clip["url"]}">{clip.get("title", "Untitled")}</a><br>\n' if clip.get("url") else f"{clip.get('title', 'Untitled')}<br>\n")
            + f"By {byline}<br>\n{clip.get('date', '')}</p>\n"
        )

        body_text = clip.get("extracted_text", "")
        body_paragraphs = _prepare_media_clip_body(body_text, clip.get("title", ""))
        email_body_text = "\n\n".join(body_paragraphs)
        if body_paragraphs:
            for para in body_paragraphs:
                para = para.strip()
                if not para:
                    continue
                body_p = doc.add_paragraph(para)
                body_p.paragraph_format.space_after = Pt(12)
                for run in body_p.runs:
                    _set_font(run)
                html_body += f"<p>{escape(para)}</p>\n"
        else:
            placeholder = doc.add_paragraph("[PASTE FULL TEXT HERE]")
            for run in placeholder.runs:
                _set_font(run)

        email_body += (
            f"{clip.get('source', 'Unknown')}\n{clip.get('title', 'Untitled')}\n{clip.get('url', '')}\n"
            f"By {byline}\n{clip.get('date', '')}\n\n{email_body_text}\n\n----------\n\n"
        )
        html_body += "<hr style='border:none;border-top:1px dashed #ccc;'>\n"

    html_body += "<p>Best regards</p>\n</div>"

    stamp = _media_clips_date_stamp(report_date_str)
    docx_path = output_dir / f"media_clips_{stamp}.docx"
    email_txt = output_dir / f"media_clips_{stamp}_email.txt"
    email_html = output_dir / f"media_clips_{stamp}_email.html"
    doc.save(docx_path)
    email_txt.write_text(email_body, encoding="utf-8")
    email_html.write_text(html_body, encoding="utf-8")
    return docx_path, email_txt, email_html


def _handle_hearing_memo(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    update_job(job_id, status="processing", progress=5, message="Preparing hearing memo inputs...")

    upload_paths = _save_uploads(output_dir, uploads, "file")
    source_path: Path | None = upload_paths[0] if upload_paths else None
    youtube_url = (values.get("youtube_url") or [""])[0].strip()

    if not source_path and youtube_url:
        transcript_text = _fetch_youtube_transcript(youtube_url)
        source_path = output_dir / "youtube_transcript.txt"
        source_path.write_text(transcript_text, encoding="utf-8")

    if not source_path:
        raise RuntimeError("No input provided. Upload a transcript or provide a YouTube URL.")

    docx_path = output_dir / "hearing_memo.docx"
    json_path = output_dir / "hearing_memo.json"
    text_path = output_dir / "hearing_memo.txt"

    cmd = [
        sys.executable,
        str(TOOLS_ROOT / "hearing_memo" / "execution" / "run.py"),
        "--input",
        str(source_path),
        "--output",
        str(docx_path),
        "--json-output",
        str(json_path),
        "--text-output",
        str(text_path),
    ]

    flag_map = {
        "memo_from": "--from",
        "memo_date": "--memo-date",
        "subject_override": "--subject",
        "hearing_title": "--hearing-title",
        "hearing_date": "--hearing-date",
        "hearing_time": "--hearing-time",
        "committee": "--committee",
        "confidentiality": "--confidentiality-footer",
    }
    for field, flag in flag_map.items():
        value = (values.get(field) or [""])[0].strip()
        if value:
            cmd.extend([flag, value])

    proc = _run_command(
        job_id,
        cmd,
        TOOLS_ROOT / "hearing_memo",
        "Running hearing memo pipeline...",
        progress=20,
        timeout=600,
        llm_model=llm_model,
    )

    verification_path = output_dir / "hearing_memo_verification.json"
    output_json = _read_json(json_path, default={}) or {}
    verification = _read_json(verification_path, default=output_json.get("verification", {})) or {}
    memo_text = _read_text(text_path)

    # Rename DOCX to <YYYYMMDD>_<title_slug>.docx
    import re as _re
    _meta = (output_json.get("record") or {}).get("metadata") or output_json.get("metadata") or {}
    _raw_date = (
        _meta.get("hearing_date")
        or (values.get("hearing_date") or [""])[0].strip()
        or (values.get("memo_date") or [""])[0].strip()
        or date.today().isoformat()
    )
    try:
        from datetime import datetime as _dt
        for _fmt in ("%B %d, %Y", "%b %d, %Y", "%Y-%m-%d", "%m/%d/%Y"):
            try:
                _date_str = _dt.strptime(_raw_date.strip(), _fmt).strftime("%Y%m%d")
                break
            except ValueError:
                continue
        else:
            _date_str = _re.sub(r"[^0-9]", "", _raw_date)[:8] or date.today().strftime("%Y%m%d")
    except Exception:
        _date_str = date.today().strftime("%Y%m%d")

    _raw_title = (
        _meta.get("hearing_title")
        or (values.get("hearing_title") or [""])[0].strip()
        or (values.get("subject_override") or [""])[0].strip()
        or "hearing_memo"
    )
    _title_slug = _re.sub(r"[^a-z0-9]+", "_", _raw_title.lower()).strip("_")[:60]
    _named_docx = output_dir / f"{_date_str}_{_title_slug}.docx"
    if docx_path.exists():
        docx_path.rename(_named_docx)
        docx_path = _named_docx

    ordered_artifacts = _existing([docx_path, text_path, verification_path, json_path])
    set_job_artifacts(job_id, ordered_artifacts)
    update_job(
        job_id,
        status="completed",
        progress=100,
        message="Hearing memo generated.",
        result_data={
            "memo_text": memo_text,
            "verdict": verification.get("verdict", ""),
            "flags": verification.get("flags", []),
            "human_checks": verification.get("human_checks", []),
            "verification": verification,
            # Pipeline logs — always surface so the user can diagnose LLM errors
            "stdout": (proc.stdout or "").strip(),
            "stderr": (proc.stderr or "").strip(),
        },
    )


def _handle_influence_tracker(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    del uploads
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    result = _run_disclosure_tracker(
        job_id,
        output_dir,
        entities=(values.get("entities") or [""])[0].strip(),
        search_field=(values.get("search_field") or ["both"])[0].strip() or "both",
        filing_years=(values.get("filing_years") or [""])[0].strip() or None,
        all_years=((values.get("all_years") or ["false"])[0].strip().lower() == "true"),
        quarters=(values.get("quarters") or ["Q1,Q2,Q3,Q4"])[0].strip() or "Q1,Q2,Q3,Q4",
        from_date=(values.get("from_date") or values.get("from") or [""])[0].strip() or None,
        to_date=(values.get("to_date") or values.get("to") or [""])[0].strip() or None,
        sources=(values.get("sources") or ["lda,irs990"])[0].strip() or "lda,irs990",
        mode=(values.get("mode") or ["basic"])[0].strip() or "basic",
        max_results=(values.get("max_results") or ["500"])[0].strip() or "500",
        fuzzy_threshold=(values.get("fuzzy_threshold") or ["85"])[0].strip() or "85",
        dry_run=((values.get("dry_run") or ["false"])[0].strip().lower() == "true"),
        llm_model=llm_model,
    )

    set_job_artifacts(job_id, result.pop("artifacts"))
    update_job(
        job_id,
        status="completed",
        progress=100,
        message="Disclosure search complete.",
        result_data=result,
    )


def _handle_legislative_tracker(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    del uploads
    output_dir = _job_output_dir(job_id)
    exec_dir = TOOLS_ROOT / "legislative_tracker" / "execution"
    _ensure_sys_path(exec_dir)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"

    action = (values.get("action") or ["search"])[0].strip() or "search"
    bill_id = (values.get("bill_id") or [""])[0].strip()
    state = (values.get("state") or ["US"])[0].strip() or "US"
    year = (values.get("year") or [""])[0].strip()
    query = (values.get("query") or [""])[0].strip()
    max_results = (values.get("max_results") or [""])[0].strip()
    title_only = (values.get("title_only") or ["false"])[0].strip().lower() == "true"
    summary_level = (values.get("summary_level") or ["preview"])[0].strip().lower() or "preview"

    if action == "search":
        if not query:
            raise RuntimeError("Query is required.")
        cmd = [
            sys.executable,
            str(exec_dir / "run.py"),
            "--json",
            "--query",
            query,
            "--state",
            state,
            "--out",
            str(output_dir),
        ]
        if year:
            cmd.extend(["--year", year])
        if max_results:
            cmd.extend(["--max-results", max_results])
        if title_only:
            cmd.append("--title-only")
        result = _run_command(job_id, cmd, exec_dir, "Searching legislation...", progress=20, timeout=180, llm_model=llm_model)
        results_json = _extract_json_stdout(result.stdout)
        if results_json is None:
            results_json = _read_json(output_dir / "search_results.json", default=[])
        report = _read_text(output_dir / "report.md")
        artifacts = _existing([output_dir / "report.md", output_dir / "search_results.json"])
        set_job_artifacts(job_id, artifacts)
        update_job(
            job_id,
            status="completed",
            progress=100,
            message="Legislative search complete.",
            result_data={"action": "search", "results": results_json, "report": report},
        )
        return

    if action == "summarize":
        if not bill_id:
            raise RuntimeError("bill_id is required for summarize.")
        cmd = [
            sys.executable,
            str(exec_dir / "run.py"),
            "--json",
            "--bill-id",
            bill_id,
            "--summarize",
            "--summary-level",
            summary_level,
            "--out",
            str(output_dir),
            "--model",
            llm_model,
        ]
        timeout = 240 if summary_level == "preview" else 3600
        result = _run_command(
            job_id,
            cmd,
            exec_dir,
            "Generating bill preview..." if summary_level == "preview" else "Generating detailed bill summary...",
            progress=25,
            timeout=timeout,
            llm_model=llm_model,
        )
        summary_json = _extract_json_stdout(result.stdout) or {}
        summary_md = _read_text(next(output_dir.rglob("bill_summary.md"), output_dir / "bill_summary.md"))
        detail_json_path = next(output_dir.rglob("bill_detail.json"), output_dir / "bill_detail.json")
        artifacts = _existing([detail_json_path, output_dir / "bill_summary.md"])
        set_job_artifacts(job_id, artifacts)
        summary_status = summary_json.get("summary_status") or "failed_system"
        source_text_status = summary_json.get("source_text_status") or "missing"
        source_status = summary_json.get("source_status") or "unusable_text"
        extraction_status = summary_json.get("extraction_status") or "not_run"
        verification_status = summary_json.get("verification_status") or "failed_system"
        extraction_coverage = summary_json.get("extraction_coverage")
        coverage_mode = summary_json.get("coverage_mode") or "metadata_only"
        evidence_coverage = summary_json.get("evidence_coverage") or extraction_coverage
        validation_flags = summary_json.get("validation_flags") or []
        unsupported_claims = summary_json.get("unsupported_claims") or []
        traceability_report = summary_json.get("traceability_report") or []
        model_path = summary_json.get("model_path") or {}
        evidence_index = summary_json.get("evidence_index") or []
        summary_structured = summary_json.get("summary_structured") or {}
        summary_level = summary_json.get("summary_level") or summary_level
        report_markdown = summary_json.get("report_markdown") or summary_md
        result_status = "completed" if summary_status in {"preview_ready", "verified", "blocked_missing_source", "blocked_verification"} else "failed"
        result_message = (
            "Bill preview generated."
            if summary_status == "preview_ready"
            else
            "Verified bill summary generated."
            if summary_status == "verified"
            else "Verified summary unavailable: official bill text was missing or unusable."
            if summary_status == "blocked_missing_source"
            else "Verified summary unavailable: generated claims could not be fully traced to the bill text."
            if summary_status == "blocked_verification"
            else "Bill summary generation failed."
        )
        update_job(
            job_id,
            status=result_status,
            progress=100,
            message=result_message,
            result_data={
                "action": "summarize",
                "bill": summary_json.get("bill") or _read_json(detail_json_path, default={}),
                "summary": summary_json.get("summary") or summary_md,
                "caveats": summary_json.get("caveats") or [],
                "summary_status": summary_status,
                "source_text_status": source_text_status,
                "source_status": source_status,
                "extraction_status": extraction_status,
                "verification_status": verification_status,
                "extraction_coverage": extraction_coverage,
                "coverage_mode": coverage_mode,
                "evidence_coverage": evidence_coverage,
                "validation_flags": validation_flags,
                "unsupported_claims": unsupported_claims,
                "traceability_report": traceability_report,
                "model_path": model_path,
                "evidence_index": evidence_index,
                "summary_structured": summary_structured,
                "report_markdown": report_markdown,
                "summary_level": summary_level,
            },
        )
        return

    legiscan_mod = _load_module("legiscan_client", exec_dir / "legiscan_client.py")
    LegiScanClient = legiscan_mod.LegiScanClient
    report_mod = _load_module("report", exec_dir / "report.py")
    ReportGenerator = report_mod.ReportGenerator
    watchlist_mod = _load_module("watchlist", exec_dir / "watchlist.py")
    WatchlistManager = watchlist_mod.WatchlistManager

    LEGISLATIVE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    watchlist = WatchlistManager(watchlist_path=str(LEGISLATIVE_CACHE_DIR / "watchlist.json"))
    client = None if action in {"watchlist_list", "watchlist_remove"} else LegiScanClient(cache_dir=str(LEGISLATIVE_CACHE_DIR))

    if action == "watchlist_add":
        if not bill_id:
            raise RuntimeError("bill_id is required for watchlist_add.")
        bill = client.get_bill(int(bill_id))
        added = watchlist.add(int(bill_id), bill)
        bills = watchlist.list_bills()
        report = ReportGenerator.watchlist_report(bills)
        report_path = output_dir / "watchlist_report.md"
        report_path.write_text(report, encoding="utf-8")
        set_job_artifacts(job_id, _existing([report_path, LEGISLATIVE_CACHE_DIR / "watchlist.json"]))
        update_job(job_id, status="completed", progress=100, message="Watchlist updated.", result_data={"action": action, "added": added, "bill": bill, "watchlist": bills, "report": report})
        return

    if action == "watchlist_remove":
        if not bill_id:
            raise RuntimeError("bill_id is required for watchlist_remove.")
        removed = watchlist.remove(int(bill_id))
        bills = watchlist.list_bills()
        report = ReportGenerator.watchlist_report(bills)
        report_path = output_dir / "watchlist_report.md"
        report_path.write_text(report, encoding="utf-8")
        set_job_artifacts(job_id, _existing([report_path, LEGISLATIVE_CACHE_DIR / "watchlist.json"]))
        update_job(job_id, status="completed", progress=100, message="Watchlist updated.", result_data={"action": action, "removed": removed, "watchlist": bills, "report": report})
        return

    if action == "watchlist_list":
        bills = watchlist.list_bills()
        report = ReportGenerator.watchlist_report(bills)
        report_path = output_dir / "watchlist_report.md"
        report_path.write_text(report, encoding="utf-8")
        set_job_artifacts(job_id, _existing([report_path, LEGISLATIVE_CACHE_DIR / "watchlist.json"]))
        update_job(job_id, status="completed", progress=100, message="Watchlist loaded.", result_data={"action": action, "watchlist": bills, "report": report})
        return

    if action == "watchlist_refresh":
        refresh_results = watchlist.refresh_all(client)
        bills = watchlist.list_bills()
        report = ReportGenerator.watchlist_report(bills, refresh_results=refresh_results)
        report_path = output_dir / "watchlist_report.md"
        report_path.write_text(report, encoding="utf-8")
        set_job_artifacts(job_id, _existing([report_path, LEGISLATIVE_CACHE_DIR / "watchlist.json"]))
        update_job(job_id, status="completed", progress=100, message="Watchlist refreshed.", result_data={"action": action, "watchlist": bills, "refresh_results": refresh_results, "report": report})
        return

    raise RuntimeError(f"Unsupported legislative_tracker action: {action}")


def _handle_messaging_matrix(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    update_job(job_id, status="processing", progress=10, message="Building messaging matrix...")

    upload_paths = _save_uploads(output_dir, uploads)
    context_text = (values.get("context") or [""])[0]
    pasted_context = (values.get("context_paste") or [""])[0]
    context_parts = [part for part in [context_text, pasted_context, _read_uploaded_context(upload_paths)] if part]
    combined_context = "\n\n".join(context_parts)

    generator = _load_module("mm_generator_fastapi", TOOLS_ROOT / "messaging_matrix" / "execution" / "generator.py")
    exporter = _load_module("mm_export_fastapi", TOOLS_ROOT / "messaging_matrix" / "execution" / "export.py")

    variants_raw = (values.get("variants") or [""])[0].strip()
    variants = [variant.strip() for variant in variants_raw.split(",") if variant.strip()] or None

    update_job(job_id, status="processing", progress=20, message="Generating message house (step 1/2)…")
    with _llm_env(llm_model):
        result = generator.generate_matrix(
            position=(values.get("position") or [""])[0],
            context=combined_context,
            organization=(values.get("organization") or [""])[0],
            target_audience=(values.get("target_audience") or [""])[0],
            core_messages=(values.get("core_messages") or [""])[0],
            facts=(values.get("facts") or [""])[0],
            variants=variants,
            style_guides_dir=str(STYLE_GUIDES_DIR),
        )
    update_job(job_id, status="processing", progress=85, message="Exporting deliverables (step 2/2)…")
    markdown = generator.render_markdown(result)

    md_path = output_dir / "messaging_matrix.md"
    json_path = output_dir / "messaging_matrix.json"
    docx_path = output_dir / "messaging_matrix.docx"
    md_path.write_text(markdown, encoding="utf-8")
    json_path.write_text(json.dumps(result, indent=2), encoding="utf-8")
    exporter.export_docx(result, str(docx_path))

    set_job_artifacts(job_id, _existing([docx_path, md_path, json_path]))
    update_job(
        job_id,
        status="completed",
        progress=100,
        message="Messaging matrix generated.",
        result_data={**result, "markdown": markdown},
    )


def _handle_stakeholder_briefing(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    update_job(job_id, status="processing", progress=10, message="Generating stakeholder briefing...")

    upload_paths = _save_uploads(output_dir, uploads)
    context_text = (values.get("context") or [""])[0]
    combined_context = "\n\n".join(part for part in [context_text, _read_uploaded_context(upload_paths)] if part)

    generator = _load_module("sb_generator_fastapi", TOOLS_ROOT / "stakeholder_briefing" / "execution" / "generator.py")
    exporter = _load_module("sb_export_fastapi", TOOLS_ROOT / "stakeholder_briefing" / "execution" / "export.py")

    include_disclosures = (values.get("include_disclosures") or ["true"])[0].strip().lower() != "false"
    include_news = (values.get("include_news") or ["true"])[0].strip().lower() != "false"

    with _llm_env(llm_model):
        result = generator.generate_briefing(
            stakeholder_name=(values.get("stakeholder_name") or values.get("stakeholder") or values.get("name") or [""])[0],
            meeting_purpose=(values.get("meeting_purpose") or values.get("purpose") or [""])[0],
            organization=(values.get("organization") or [""])[0],
            your_organization=(values.get("your_org") or [""])[0],
            context=combined_context,
            include_disclosures=include_disclosures,
            include_news=include_news,
        )
    markdown = generator.render_markdown(result)

    md_path = output_dir / "stakeholder_briefing.md"
    json_path = output_dir / "stakeholder_briefing.json"
    docx_path = output_dir / "stakeholder_briefing.docx"
    md_path.write_text(markdown, encoding="utf-8")
    json_path.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    exporter.export_docx(result, str(docx_path))

    set_job_artifacts(job_id, _existing([docx_path, md_path, json_path]))
    update_job(job_id, status="completed", progress=100, message="Stakeholder briefing generated.", result_data={**result, "markdown": markdown})


def _handle_media_list_builder(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    del uploads
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    update_job(job_id, status="processing", progress=10, message="Building media list...")

    generator = _load_module("mlb_generator_fastapi", TOOLS_ROOT / "media_list_builder" / "execution" / "generator.py")
    exporter = _load_module("mlb_export_fastapi", TOOLS_ROOT / "media_list_builder" / "execution" / "export.py")

    media_types_raw = (values.get("media_types") or values.get("outlet_types") or [""])[0].strip()
    media_types = [item.strip() for item in media_types_raw.split(",") if item.strip()] or None
    location = (values.get("location") or values.get("region") or ["US"])[0].strip() or "US"

    source_filter = (values.get("source_filter") or ["national"])[0].strip() or "national"

    with _llm_env(llm_model):
        result = generator.generate_media_list(
            issue=(values.get("issue") or values.get("topic") or [""])[0],
            location=location,
            media_types=media_types,
            num_contacts=int((values.get("num_contacts") or ["20"])[0] or "20"),
            source_filter=source_filter,
            broad_topic=(values.get("broad_topic") or [""])[0].strip(),
            coverage_desk=(values.get("coverage_desk") or [""])[0].strip(),
            topic_mode=(values.get("topic_mode") or ["specific"])[0].strip() or "specific",
        )
    markdown = generator.render_markdown(result)

    md_path = output_dir / "media_list.md"
    json_path = output_dir / "media_list.json"
    xlsx_path = output_dir / "media_list.xlsx"
    md_path.write_text(markdown, encoding="utf-8")
    json_path.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    exporter.export_xlsx(result, str(xlsx_path))

    set_job_artifacts(job_id, _existing([xlsx_path, md_path, json_path]))
    update_job(job_id, status="completed", progress=100, message="Media list generated.", result_data={**result, "markdown": markdown})


def _handle_stakeholder_map(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    del uploads
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    update_job(job_id, status="processing", progress=10, message="Building stakeholder map...")

    generator = _load_module("smb_generator_fastapi", TOOLS_ROOT / "stakeholder_map" / "execution" / "generator.py")
    exporter = _load_module("smb_export_fastapi", TOOLS_ROOT / "stakeholder_map" / "execution" / "export.py")
    analytics_mod = _load_module("smb_analytics_fastapi", TOOLS_ROOT / "stakeholder_map" / "execution" / "analytics.py")

    scope = (values.get("scope") or ["federal"])[0].strip().lower() or "federal"
    state = (values.get("state") or ["US"])[0].strip().upper() or "US"
    include_types_raw = (values.get("include_types") or [""])[0].strip()
    include_types = [item.strip().lower() for item in include_types_raw.split(",") if item.strip()] or None

    with _llm_env(llm_model):
        result = generator.build_map(
            policy_issue=(values.get("policy_issue") or [""])[0],
            scope=scope,
            state=state,
            include_types=include_types,
        )
    markdown = generator.render_markdown(result)
    # analytics is computed inside build_map() and enriches actors in-place;
    # fall back to a fresh compute if missing (e.g. older cached result)
    analytics = result.get("network_analytics") or analytics_mod.compute_network_analytics(
        result.get("actors", []), result.get("relationships", [])
    )

    md_path = output_dir / "stakeholder_map.md"
    json_path = output_dir / "stakeholder_map.json"
    xlsx_path = output_dir / "stakeholder_map.xlsx"
    docx_path = output_dir / "stakeholder_map.docx"
    html_path = output_dir / "stakeholder_map.html"

    md_path.write_text(markdown, encoding="utf-8")
    json_path.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    exporter.export_xlsx(result, str(xlsx_path))
    exporter.export_docx(result, str(docx_path))

    try:
        graph_mod = _load_module("smb_graph_fastapi", TOOLS_ROOT / "stakeholder_map" / "execution" / "graph.py")
        fig = graph_mod.build_network_graph(
            actors=result.get("actors", []),
            relationships=result.get("relationships", []),
            title=f"Stakeholder Map: {(values.get('policy_issue') or [''])[0]}",
            centrality={actor["id"]: actor.get("betweenness_centrality", 0) for actor in result.get("actors", [])},
        )
        graph_mod.save_graph_html(fig, str(html_path))
    except Exception:
        pass

    set_job_artifacts(job_id, _existing([xlsx_path, docx_path, json_path, md_path, html_path]))
    update_job(
        job_id,
        status="completed",
        progress=100,
        message="Stakeholder map generated.",
        result_data={
            **result,
            "markdown": markdown,
            "analytics": analytics,
            "strategic_analysis": result.get("strategic_analysis", {}),
        },
    )


def _handle_background_memo(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    update_job(job_id, status="processing", progress=5, message="Preparing background memo research...")

    upload_paths = _save_uploads(output_dir, uploads)
    file_context = _read_uploaded_context(upload_paths)

    subject = (values.get("subject") or [""])[0].strip()
    memo_date = (values.get("memo_date") or [date.today().strftime("%B %d, %Y")])[0].strip() or date.today().strftime("%B %d, %Y")
    context = (values.get("context") or [""])[0]
    sections_raw = (values.get("sections_text") or values.get("sections") or [""])[0]
    sections = [line.strip() for line in re.split(r"[\n,]+", sections_raw) if line.strip()]
    if not sections:
        raise RuntimeError("At least one section heading is required.")

    disclosure_entity = (values.get("disclosure_entity_override") or [""])[0].strip() or subject
    disclosure_sources = (values.get("disclosure_sources") or [""])[0].strip()

    research_md = ""
    disclosure_md = ""
    disclosure_csv_data: dict[str, Any] = {}
    disclosure_artifacts: list[Path] = []

    research_path = TOOLS_ROOT / "background_memo" / "execution" / "research.py"
    if research_path.exists():
        try:
            update_job(job_id, status="processing", progress=20, message="Running background web research...")
            research_mod = _load_module("bmg_research_fastapi", research_path)
            research_md = research_mod.research_subject(subject, context) or ""
        except Exception:
            research_md = ""

    if disclosure_sources:
        try:
            update_job(job_id, status="processing", progress=35, message="Searching disclosure records...")
            disclosure_result = _run_disclosure_tracker(
                job_id,
                output_dir / "disclosures",
                entities=disclosure_entity,
                search_field="both",
                all_years=True,
                quarters="Q1,Q2,Q3,Q4",
                sources=disclosure_sources,
                mode="basic",
            )
            disclosure_md = disclosure_result["report"]
            disclosure_csv_data = disclosure_result["csv_data"]
            disclosure_artifacts = [Path(path) for path in disclosure_result.get("artifacts", [])]
        except Exception:
            disclosure_md = ""
            disclosure_csv_data = {}
            disclosure_artifacts = []

    generator = _load_module("bmg_generator_fastapi", TOOLS_ROOT / "background_memo" / "execution" / "generator.py")
    exporter = _load_module("bmg_export_fastapi", TOOLS_ROOT / "background_memo" / "execution" / "export.py")

    update_job(job_id, status="processing", progress=60, message="Generating background memo...")
    combined_research = "\n\n".join(part for part in [file_context, research_md] if part)
    with _llm_env(llm_model):
        result = generator.generate_memo(
            subject=subject,
            sections=sections,
            context=context,
            disclosure_context=disclosure_md,
            research_context=combined_research,
            suppress_disclosures=not bool(disclosure_sources),
        )
    markdown = generator.render_markdown(result, memo_date=memo_date)

    docx_path = output_dir / "background_memo.docx"
    md_path = output_dir / "background_memo.md"
    json_path = output_dir / "background_memo.json"
    md_path.write_text(markdown, encoding="utf-8")
    json_path.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    exporter.export_docx(result, str(docx_path), memo_date=memo_date)

    artifacts = [docx_path, md_path, json_path, *disclosure_artifacts]
    set_job_artifacts(job_id, _existing(artifacts))

    # Build a validated, flat result_data using the job-level schema.
    schema_mod = _load_module("bmg_schema_fastapi", BMG_SCHEMA_PATH)
    job_result = schema_mod.BackgroundMemoJobResult(
        subject=result["subject"],
        memo_date=memo_date,
        sections_requested=result.get("sections_requested", sections),
        overview=result["overview"],
        fast_facts=result["fast_facts"],
        sections=result["sections"],
        links=result["links"],
        markdown=markdown,
        research_md=research_md,
        disclosure_md=disclosure_md,
        disclosure_csv_data=disclosure_csv_data,
    )
    update_job(
        job_id,
        status="completed",
        progress=100,
        message="Background memo generated.",
        result_data=job_result.model_dump(),
    )


def _handle_media_clips(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    del uploads
    output_dir = _job_output_dir(job_id)
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    action = (values.get("action") or ["generate"])[0].strip() or "generate"

    if action == "build_report":
        update_job(job_id, status="processing", progress=15, message="Rebuilding clips report...")
        clips_data_raw = (values.get("clips_data_json") or ["[]"])[0]
        clips_data = json.loads(clips_data_raw)
        report_topic = (values.get("report_topic") or values.get("topic") or ["Media Clips"])[0]
        report_date = (values.get("report_date_str") or values.get("target_date") or [date.today().isoformat()])[0]
        docx_path, email_txt, email_html = _build_media_report(clips_data, report_topic, report_date, output_dir)
        json_path = output_dir / "clips_data.json"
        json_path.write_text(json.dumps(clips_data, indent=2), encoding="utf-8")
        set_job_artifacts(job_id, _existing([docx_path, email_txt, email_html, json_path]))
        email_body_text = _read_text(email_txt) if email_txt.exists() else ""
        update_job(job_id, status="completed", progress=100, message="Media clips report rebuilt.", result_data={"clips_data": clips_data, "missing_full_text": sum(1 for clip in clips_data if not clip.get("has_full_text")), "email_body": email_body_text})
        return

    cmd = [
        sys.executable,
        str(TOOLS_ROOT / "media_clips" / "execution" / "generate_clips.py"),
        "--topic",
        (values.get("topic") or [""])[0],
        "--queries",
        ",".join(
            line.strip()
            for line in (values.get("queries") or [""])[0].splitlines()
            if line.strip()
        )
        or (values.get("queries") or [""])[0],
        "--period",
        (values.get("period") or ["24h"])[0],
        "--output-dir",
        str(output_dir),
        "--no-email",
    ]
    since = (values.get("since") or values.get("since_date") or [""])[0].strip()
    period_str = (values.get("period") or ["24h"])[0].strip()

    # Auto-compute --since from period when not manually provided.
    # GNews ignores its own period parameter and returns stale articles,
    # so we enforce the cutoff ourselves via post-fetch date filtering.
    if not since:
        period_map = {"h": "hours", "d": "days", "w": "weeks", "m": "days"}
        m = re.match(r"^(\d+)([hdwm])$", period_str.lower())
        if m:
            amount = int(m.group(1))
            unit_char = m.group(2)
            if unit_char == "m":
                amount *= 30  # approximate months as days
            kwargs = {period_map[unit_char]: amount}
            since = (datetime.utcnow() - timedelta(**kwargs)).strftime("%Y-%m-%d %H:%M")

    target_date = (values.get("target_date") or [date.today().isoformat()])[0].strip()
    source_filter = (values.get("source_filter") or ["mainstream"])[0].strip().lower()
    custom_sources = (values.get("custom_sources") or [""])[0].strip()
    if since:
        cmd.extend(["--since", since])
    if target_date:
        cmd.extend(["--target-date", target_date])
    if source_filter in {"all sources", "all", "all_sources"}:
        cmd.append("--all-sources")
    elif custom_sources:
        cmd.extend(["--custom-sources", custom_sources.replace("\n", ",")])
    max_clips_str = (values.get("max_clips") or [""])[0].strip()
    if max_clips_str and max_clips_str.lower() not in {"all", "0", ""}:
        try:
            cmd.extend(["--max-clips", str(int(max_clips_str))])
        except ValueError:
            pass

    command_result = _run_command(job_id, cmd, TOOLS_ROOT / "media_clips" / "execution", "Generating media clips...", progress=20, timeout=240, llm_model=llm_model)

    data_path = next(output_dir.rglob("*_data.json"), None)
    docx_path = next(output_dir.rglob("*.docx"), None)
    clips_data = _read_json(data_path, default=[]) if data_path else []
    artifacts = _existing(([docx_path] if docx_path else []) + ([data_path] if data_path else []))
    set_job_artifacts(job_id, artifacts)
    update_job(
        job_id,
        status="completed",
        progress=100,
        message="Media clips generated.",
        result_data={
            "clips_data": clips_data,
            "missing_full_text": sum(1 for clip in clips_data if not clip.get("has_full_text")),
            "stdout": command_result.stdout or "",
            "stderr": command_result.stderr or "",
        },
    )


def _handle_media_clip_cleaner(job_id: str, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    output_dir = _job_output_dir(job_id)
    update_job(job_id, status="processing", progress=10, message="Cleaning article text...")

    upload_paths = _save_uploads(output_dir, uploads)
    output_path = output_dir / "cleaned_clip.md"
    raw_text = (values.get("raw_text") or [""])[0]
    title = (values.get("title") or [""])[0].strip()
    mode = (values.get("mode") or ["local"])[0].strip().lower() or "local"
    llm_model = (values.get("llm_model") or ["ChangeAgent"])[0].strip() or "ChangeAgent"
    cmd = [
        sys.executable,
        str(TOOLS_ROOT / "media_clip_cleaner" / "execution" / "run.py"),
        "--output-file",
        str(output_path),
        "--mode",
        "llm" if mode == "llm" else "local",
    ]
    if title:
        cmd.extend(["--title", title])
    if mode == "llm" and llm_model:
        cmd.extend(["--llm-model", llm_model])
    if (values.get("fallback_local") or ["false"])[0].strip().lower() == "true":
        cmd.append("--fallback-local")
    if raw_text.strip():
        cmd.extend(["--raw-text", raw_text])
    elif upload_paths:
        cmd.extend(["--input-file", str(upload_paths[0])])
    else:
        raise RuntimeError("Provide raw_text or an uploaded file.")

    _run_command(job_id, cmd, TOOLS_ROOT / "media_clip_cleaner" / "execution", "Running clip cleaner...", progress=25, timeout=180, llm_model=llm_model)

    cleaned_text = _read_text(output_path)
    set_job_artifacts(job_id, _existing([output_path]))
    update_job(job_id, status="completed", progress=100, message="Clip cleaned.", result_data={"cleaned_text": cleaned_text})


HANDLERS = {
    "hearing_memo_generator": _handle_hearing_memo,
    "influence_disclosure_tracker": _handle_influence_tracker,
    "legislative_tracker": _handle_legislative_tracker,
    "messaging_matrix": _handle_messaging_matrix,
    "stakeholder_briefing": _handle_stakeholder_briefing,
    "media_list_builder": _handle_media_list_builder,
    "stakeholder_map": _handle_stakeholder_map,
    "background_memo_generator": _handle_background_memo,
    "media_clips": _handle_media_clips,
    "media_clip_cleaner": _handle_media_clip_cleaner,
}


def _run_handler(job_id: str, handler, values: dict[str, list[str]], uploads: list[dict[str, Any]]) -> None:
    try:
        handler(job_id, values, uploads)
    except Exception as exc:
        message = str(exc)
        lowered = message.lower()
        if "429" in message or "insufficient_quota" in lowered or "rate limit" in lowered:
            message = (
                "This tool is temporarily unavailable because the upstream OpenAI quota or rate limit was exceeded. "
                "Try again later or reduce concurrent runs."
            )
        update_job(job_id, status="failed", progress=100, message=message)


@router.post("/open-email-draft")
async def open_email_draft(request: Request):
    """Open a Mail.app draft pre-filled with the media clips email body and DOCX attachment."""
    form = await request.form()
    job_id = (form.get("job_id") or "").strip()
    to_raw = (form.get("to") or "").strip()
    subject = (form.get("subject") or "Media Clips").strip()
    sender = (form.get("sender") or "").strip()

    if not job_id:
        raise HTTPException(status_code=400, detail="job_id is required.")
    if not to_raw:
        raise HTTPException(status_code=400, detail="At least one recipient (to) is required.")

    output_dir = JOBS_OUTPUT_ROOT / job_id
    if not output_dir.exists():
        raise HTTPException(status_code=404, detail=f"No output directory for job {job_id}.")

    # Find the DOCX artifact
    docx_files = sorted(output_dir.glob("*.docx"))
    docx_path = str(docx_files[0].resolve()) if docx_files else None

    # Find the plain-text email body
    txt_files = sorted(output_dir.glob("*email*.txt"))
    body_text = ""
    if txt_files:
        body_text = txt_files[0].read_text(encoding="utf-8", errors="replace")
    if not body_text:
        body_text = f"Please find attached the {subject}."

    recipients = [r.strip() for r in to_raw.split(",") if r.strip()]

    # Build AppleScript
    subject_safe = subject.replace('"', '\\"')
    body_safe = body_text.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n")
    sender_prop = f'sender:"{sender}", ' if sender else ""
    recipient_list_as = ", ".join(f'"{r}"' for r in recipients)

    script = f'''
set recipientList to {{{recipient_list_as}}}
set theSubject to "{subject_safe}"
set theBody to "{body_safe}"
{"set theFile to POSIX file " + chr(34) + docx_path.replace(chr(34), chr(92)+chr(34)) + chr(34) if docx_path else ""}

tell application "Mail"
    set newMessage to make new outgoing message with properties {{{sender_prop}subject:theSubject, content:theBody, visible:true}}
    tell newMessage
        {"make new attachment with properties {file name:theFile} at after the last paragraph" if docx_path else ""}
        repeat with r in recipientList
            make new to recipient at end of to recipients with properties {{address:r}}
        end repeat
    end tell
    activate
end tell
'''

    try:
        result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True, timeout=15)
        if result.returncode != 0:
            raise HTTPException(status_code=500, detail=f"AppleScript error: {result.stderr.strip()}")
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="AppleScript timed out.")
    except FileNotFoundError:
        raise HTTPException(status_code=500, detail="osascript not found — email draft requires macOS.")

    return {"status": "ok", "message": "Mail.app draft opened."}


@router.post("/pitch-draft")
async def generate_pitch_draft(request: Request):
    """
    Generate a tailored pitch email for a single journalist contact.
    Synchronous — returns the draft immediately (no job queue).
    """
    import os
    from openai import OpenAI

    body = await request.json()
    contact: dict = body.get("contact", {})
    issue: str = body.get("issue", "")
    llm_model: str = body.get("llm_model", "ChangeAgent") or "ChangeAgent"

    if not contact or not issue:
        raise HTTPException(status_code=400, detail="contact and issue are required.")

    raw_name = f"{contact.get('first_name', '')} {contact.get('last_name', '')}".strip()
    unresolved_author = contact.get("contact_status") == "story_lead" or raw_name.lower() == "to verify"
    name = "" if unresolved_author else raw_name
    outlet = contact.get("outlet", "")
    role = contact.get("role", "")
    beat_notes = contact.get("notes", "")
    pitch_angle = contact.get("pitch_angle", "")
    why_now = contact.get("why_now", "")
    pitch_offer = contact.get("pitch_offer", "")
    prev_title = contact.get("previous_story_title", "")
    prev_url = contact.get("previous_story_url", "")
    pitch_root = TOOLS_ROOT / "media_list_builder"

    def _read_optional(path: Path) -> str:
        return path.read_text(encoding="utf-8").strip() if path.exists() else ""

    best_practices = _read_optional(pitch_root / "pitch_instructions" / "pitch_best_practices.txt")
    style_guide = _read_optional(pitch_root / "PITCH_STYLE_GUIDE.md")
    agent_instructions = _read_optional(pitch_root / "pitch_agent_instructions.md")
    skill_text = _read_optional(pitch_root / "skill.md")
    example_paths = sorted((pitch_root / "pitch_examples").glob("*.txt"))
    example_block = "\n\n".join(
        f"Example {idx}:\n{path.read_text(encoding='utf-8').strip()}"
        for idx, path in enumerate(example_paths, start=1)
    )

    system_prompt = (
        "You are drafting a one-to-one journalist pitch using the project's reusable pitch instruction layer.\n\n"
        "Follow the source hierarchy exactly:\n"
        "1. pitch best practices\n"
        "2. pitch style guide\n"
        "3. agent instructions\n"
        "4. examples for calibration only\n\n"
        "Examples are reference only. Do not copy their topic framing, sentence structure, cadence, or tone too closely.\n\n"
        f"=== PITCH BEST PRACTICES ===\n{best_practices}\n\n"
        f"=== PITCH STYLE GUIDE ===\n{style_guide}\n\n"
        f"=== PITCH AGENT INSTRUCTIONS ===\n{agent_instructions}\n\n"
        f"=== PITCH SKILL ===\n{skill_text}\n\n"
        f"=== CALIBRATION EXAMPLES ONLY ===\n{example_block}\n"
    )

    user_prompt = (
        f"Issue being pitched: {issue}\n\n"
        f"Journalist: {name or 'Unknown'}\n"
        f"Outlet: {outlet}\n"
        f"Role/Beat: {role}\n"
        f"Pitch angle: {pitch_angle}\n"
        f"Why-now hook: {why_now}\n"
        f"Concrete offer: {pitch_offer}\n"
        f"Recent relevant story: \"{prev_title}\" ({prev_url})\n"
        f"Additional context: {beat_notes}\n\n"
        "Draft a pitch email as JSON. Hard rules:\n"
        f"- Author resolved: {'no' if unresolved_author else 'yes'}.\n"
        "- If author resolved is no, do not guess or invent a journalist name. Write a pitch that can be sent once the correct reporter name is confirmed, and avoid personal-name greetings.\n"
        "- Subject: name the development or access, not the topic category. Under 10 words.\n"
        "- Opening line: reference the specific prior story in one sentence, then move on.\n"
        "- Development: state what specifically happened, changed, or is imminent. Name the bill, agency, date, or ruling.\n"
        "- Offer: use the 'Concrete offer' field above. One offer only. Do not invent exclusives or data.\n"
        "- CTA: one sentence, light, non-pushy.\n"
        "- Body: under 200 words total.\n"
        "- Never use: 'emerging legislative language', 'timely update', 'brief briefing', 'gaining traction', 'given recent developments'.\n\n"
        "Return:\n"
        "{\"subject\": \"...\", \"body\": \"...\"}"
    )

    with _llm_env(llm_model):
        api_key = os.environ.get("OPENAI_API_KEY", "")
        base_url = os.environ.get("OPENAI_BASE_URL")
        active_model = os.environ.get("LLM_MODEL_OVERRIDE") or llm_model

        kwargs: dict = {"api_key": api_key}
        if base_url:
            kwargs["base_url"] = base_url
        client = OpenAI(**kwargs)

        response_format_kwarg = {} if os.environ.get("LLM_MODEL_OVERRIDE") else {"response_format": {"type": "json_object"}}

        resp = client.chat.completions.create(
            model=active_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.4,
            **response_format_kwarg,
        )

    content = resp.choices[0].message.content or ""
    # Parse JSON, fall back gracefully
    import re as _re
    text = content.strip()
    m = _re.search(r"```(?:json)?\s*([\s\S]+?)```", text)
    if m:
        text = m.group(1).strip()
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        data = {"subject": f"Story pitch: {issue[:60]}", "body": text}

    return {"subject": data.get("subject", ""), "body": data.get("body", "")}


@router.get("")
def list_tools():
    """Return tool metadata from tool-registry.yaml, annotated with handler availability."""
    import yaml  # noqa: PLC0415 — lazy import keeps startup fast if pyyaml somehow missing

    registry_path = TOOLKIT_ROOT / "tool-registry.yaml"
    if not registry_path.exists():
        return []
    data = yaml.safe_load(registry_path.read_text(encoding="utf-8"))
    tools = data.get("tools", [])
    for tool in tools:
        tool["has_handler"] = tool["id"] in HANDLERS
    return tools


@router.post("/execute/{tool_id}")
async def execute_tool(tool_id: str, request: Request):
    handler = HANDLERS.get(tool_id)
    if not handler:
        raise HTTPException(status_code=404, detail=f"Unknown tool: {tool_id}")

    form = await request.form()
    values: dict[str, list[str]] = {}
    uploads: list[dict[str, Any]] = []

    for key, value in form.multi_items():
        if isinstance(value, UploadFile):
            uploads.append(
                {
                    "field": key,
                    "filename": value.filename,
                    "content": await value.read(),
                }
            )
        else:
            values.setdefault(key, []).append(str(value))

    job_id = create_job()
    thread = threading.Thread(target=_run_handler, args=(job_id, handler, values, uploads), daemon=True)
    thread.start()
    return {"job_id": job_id}
