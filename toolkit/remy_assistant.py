"""
Remy assistant backend (FastAPI).

Remy is a tool-aware public affairs assistant that can explain, recommend,
and execute the toolkit's existing tools through their CLI entry points.
"""

from __future__ import annotations

import csv
import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
import yaml

TOOLKIT_ROOT = Path(__file__).resolve().parent
REGISTRY_PATH = TOOLKIT_ROOT / "tool-registry.yaml"
OUTPUT_ROOT = TOOLKIT_ROOT / "output" / "remy"
STYLE_GUIDES_ROOT = TOOLKIT_ROOT / "tools" / "messaging_matrix" / "style_samples" / "style_guides"

_FRONTEND_PATHS: dict[str, str] = {
    "media_clips": "/media-clips",
    "media_clip_cleaner": "/media-clips",
    "influence_disclosure_tracker": "/influence",
    "hearing_memo_generator": "/memos",
    "legislative_tracker": "/legislative",
    "messaging_matrix": "/messaging",
    "media_list_builder": "/media-list",
    "stakeholder_briefing": "/stakeholder-briefing",
    "background_memo_generator": "/background-memo",
    "stakeholder_map": "/stakeholders",
}


TOOL_GUIDANCE = {
    "hearing_memo_generator": {
        "when_to_use": "Use for congressional hearing transcripts when the user needs a verified, house-style memo.",
        "notes": "Best when a transcript PDF/TXT is uploaded. Output requires human review because risk is medium.",
    },
    "media_clips": {
        "when_to_use": "Use for daily or rapid media monitoring built from news queries.",
        "notes": "Remy disables email draft creation by default and returns the report artifacts only.",
    },
    "media_clip_cleaner": {
        "when_to_use": "Use to clean copied article text into clip-ready body text.",
        "notes": "The local cleaner is the safe default; LLM cleaning can be requested explicitly.",
    },
    "influence_disclosure_tracker": {
        "when_to_use": "Use for LDA, FARA, and IRS 990 disclosure lookups on entities or stakeholders.",
        "notes": "Outputs are factual research artifacts and should still be checked before external use.",
    },
    "legislative_tracker": {
        "when_to_use": "Use to search bills, inspect a specific bill, or maintain a watchlist.",
        "notes": "If summarization is requested, Remy will use the tool's AI summary option.",
    },
    "messaging_matrix": {
        "when_to_use": "Use to turn a policy position into messaging deliverables for multiple audiences or channels.",
        "notes": "Context files can be uploaded and passed through to the tool.",
    },
    "media_list_builder": {
        "when_to_use": "Use to build a targeted journalist list around an issue and geography.",
        "notes": "This depends on live web research and contact discovery, so results should be reviewed manually.",
    },
    "stakeholder_briefing": {
        "when_to_use": "Use for meeting prep on a stakeholder, legislator, organization, or coalition.",
        "notes": "Supports disclosure/news research and optional uploaded context documents.",
    },
    "background_memo_generator": {
        "when_to_use": "Use for compact background memos on organizations, issues, clients, or individuals.",
        "notes": "Uploaded files can be merged into the text context before execution.",
    },
    "stakeholder_map": {
        "when_to_use": "Use to identify and classify the policy ecosystem around an issue.",
        "notes": "Outputs can include network graph files; large runs may take longer than the other tools.",
    },
}


def load_tool_catalog() -> list[dict[str, Any]]:
    """Load the canonical tool registry plus app-specific metadata."""
    with open(REGISTRY_PATH, encoding="utf-8") as f:
        registry = yaml.safe_load(f) or {}

    catalog = []
    for tool in registry.get("tools", []):
        tool_id = tool["id"]
        tool_root = TOOLKIT_ROOT / tool["path"]
        spec_path = tool_root / "spec.md"
        catalog.append(
            {
                **tool,
                "frontend_path": tool.get("frontend_path"),
                "spec_path": str(spec_path),
                "spec_summary": _read_spec_summary(spec_path),
                "assistant_guidance": TOOL_GUIDANCE.get(tool_id, {}),
            }
        )
    return catalog


def _load_writing_styles() -> str:
    """Load style guide files and return a condensed writing knowledge block."""
    guides = [
        ("General style", "general_style_guide.md"),
        ("Talking points", "talking_points_style_guide.md"),
        ("Press releases", "press_releases_style_guide.md"),
        ("Op-eds", "op_eds_style_guide.md"),
        ("Speeches", "speeches_style_guide.md"),
        ("Social media", "social_media_style_guide.md"),
    ]
    sections = []
    for label, filename in guides:
        path = STYLE_GUIDES_ROOT / filename
        if path.exists():
            content = path.read_text(encoding="utf-8").strip()
            sections.append(f"### {label}\n{content}")
    return "\n\n".join(sections) if sections else ""


def build_system_prompt(catalog: list[dict[str, Any]], uploaded_files: list[dict[str, Any]]) -> str:
    """Construct Remy's system prompt from registry metadata, writing styles, and current uploads."""
    # Only surface the two executable toolkit tools in the catalog block
    allowed_tool_ids = {"legislative_tracker", "influence_disclosure_tracker"}
    tool_lines = []
    for tool in catalog:
        if tool["id"] not in allowed_tool_ids:
            continue
        inputs = ", ".join(tool.get("inputs", {}).get("required", []))
        optional = ", ".join(tool.get("inputs", {}).get("optional", []))
        frontend_path = tool.get("frontend_path") or "not available"
        guidance = tool.get("assistant_guidance", {})
        tool_lines.append(
            "\n".join(
                [
                    f"- {tool['name']} ({tool['id']})",
                    f"  Required inputs: {inputs or 'none listed'}",
                    f"  Optional inputs: {optional or 'none listed'}",
                    f"  App page: {frontend_path}",
                    f"  When to use: {guidance.get('when_to_use') or tool.get('description', '')}",
                ]
            )
        )

    if uploaded_files:
        upload_lines = [
            f"- {item['name']} ({item['kind']}, {item['size_bytes']} bytes)"
            for item in uploaded_files
        ]
        upload_block = "Uploaded files available:\n" + "\n".join(upload_lines)
    else:
        upload_block = ""

    writing_styles = _load_writing_styles()
    writing_block = f"\n\n---\n## Writing Style Standards\n\n{writing_styles}" if writing_styles else ""

    return f"""
You are Remy, a senior public affairs strategist and writer.

Role:
- Act as a seasoned PA operator: sharp, controlled, discreet, strategic.
- You can research, analyze, and write professional PA documents.
- When writing, apply the style standards defined below — do not deviate from them.

Capabilities:
1. RESEARCH — search recent news (gnews_search) and the open web (brave_search) to gather facts before writing or advising.
2. LEGISLATIVE TRACKING — run the Legislative Tracker to find bills, check status, summarize legislation.
3. DISCLOSURE RESEARCH — run the Influence Disclosure Tracker to look up LDA, FARA, and IRS 990 filings.
4. WRITING — draft any PA document type (talking points, press release, op-ed, memo, speech, social media post, pitch, email) and produce a downloadable DOCX + markdown artifact via write_document.

Behavior rules:
- Always ask clarifying questions before executing if key inputs are missing. Go back and forth until you have what you need.
- For writing tasks: confirm the subject, position, audience, and any key facts before drafting. One round of questions is usually enough.
- For research tasks: use gnews_search and/or brave_search to gather current facts before writing — do not fabricate.
- For legislative or disclosure queries: run the relevant tool rather than answering from general knowledge.
- When producing a written document: ALWAYS call write_document with the full content to produce the artifact. Never output a document only as chat text — always package it via write_document so the user gets a downloadable file. The content parameter is the complete document in markdown.
- Never fabricate facts, names, filings, dates, or legislative details.
- Flag medium-risk outputs as requiring human review before external use.
- If a run fails, explain the blocker plainly and propose the next step.
- Always include source links in your responses. When gnews_search or brave_search return URLs, cite them inline as markdown links next to the claims they support. For LDA filings, link to https://lda.senate.gov/filings/public/filing/search/ when referencing lobbying disclosure data. For FARA, link to https://www.fara.gov/. For LegiScan bills, include the bill URL if available in the tool result.

Executable toolkit tools:
{chr(10).join(tool_lines)}
{upload_block}{writing_block}
""".strip()


def default_greeting() -> str:
    return (
        "I’m Remy. Tell me the objective, not the ceremony. "
        "I can route you to the right public-affairs tool, gather missing inputs, and run the toolkit when the inputs are clear."
    )


def _make_client(model: str):
    """Return the right OpenAI-compatible client for the given model name."""
    from openai import OpenAI

    if model == "ChangeAgent":
        return OpenAI(
            api_key=os.environ.get("CHANGE_AGENT_API_KEY"),
            base_url="https://runpod-proxy-956966668285.us-central1.run.app/v1/",
        )
    return OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


def _chat_change_agent(
    client,
    messages: list[dict[str, Any]],
) -> str:
    """Call ChangeAgent and return the response content."""
    response = client.chat.completions.create(
        model="ChangeAgent",
        messages=messages,
        temperature=0.4,
    )
    return response.choices[0].message.content or "No response returned."


def _parse_change_agent_tool_call(text: str) -> tuple[str, dict[str, Any]] | tuple[None, None]:
    """Parse ChangeAgent's XML-style function call. Handles two formats:
    <function=NAME>...</function>  and  <tool><name>NAME</name>...</tool>
    Returns (function_name, arguments) or (None, None)."""
    # Format 1: <function=NAME>..params..</function>
    fn_match = re.search(r"<function=([^\s>]+)>(.*?)(?:</function>|$)", text, re.DOTALL)
    if fn_match:
        fn_name = fn_match.group(1).strip()
        body = fn_match.group(2)
        return fn_name, _extract_xml_params(body)

    # Format 2: <tool><name>NAME</name>..params..</tool>
    tool_match = re.search(r"<tool>(.*?)(?:</tool>|$)", text, re.DOTALL)
    if tool_match:
        body = tool_match.group(1)
        name_match = re.search(r"<name>\s*([^\s<]+)\s*</name>", body)
        if name_match:
            fn_name = name_match.group(1).strip()
            return fn_name, _extract_xml_params(body)

    return None, None


def _extract_xml_params(body: str) -> dict[str, Any]:
    """Extract key/value pairs from a ChangeAgent parameter block."""
    args: dict[str, Any] = {}
    # Match both <parameter=KEY> and <parameter>KEY> styles
    for m in re.finditer(r"<parameter[=>]([^>\s]+)>\s*(.*?)\s*(?=</?parameter|</tool>|</function>|$)", body, re.DOTALL):
        key = m.group(1).rstrip(">").strip()
        raw = m.group(2).strip()
        try:
            args[key] = json.loads(raw)
        except (json.JSONDecodeError, ValueError):
            args[key] = raw
    return args


def _clean_change_agent_text(text: str) -> str:
    """Strip tool call blocks and fabricated tool_result blocks from ChangeAgent response."""
    cleaned = text
    cleaned = re.sub(r"\s*<function=.*?(?:</function>|$)", "", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r"\s*<tool>.*?(?:</tool>|$)", "", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r"\s*<tool_result>.*?(?:</tool_result>|$)", "", cleaned, flags=re.DOTALL)
    return cleaned.strip() or text.strip()


def chat_with_remy(
    user_message: str,
    history: list[dict[str, Any]] | None = None,
    uploaded_files: list[dict[str, Any]] | None = None,
    model: str = "gpt-4.1-mini",
) -> dict[str, Any]:
    """Run one assistant turn, allowing tool calls against the toolkit."""
    history = history or []
    uploaded_files = uploaded_files or []
    catalog = load_tool_catalog()

    fast_path = _maybe_autorun_tool(user_message, history, catalog, uploaded_files)
    if fast_path:
        return fast_path

    client = _make_client(model)

    # ChangeAgent: tool-call loop — parse XML function calls from response text
    if model == "ChangeAgent":
        messages: list[dict[str, Any]] = [
            {"role": "system", "content": build_system_prompt(catalog, uploaded_files)}
        ]
        messages.extend(_normalize_history(history))
        messages.append({"role": "user", "content": user_message})
        tool_events: list[dict[str, Any]] = []
        for _ in range(6):
            text = _chat_change_agent(client, messages)
            fn_name, fn_args = _parse_change_agent_tool_call(text)
            if not fn_name:
                return {"text": text, "tool_events": tool_events}
            result = _dispatch_tool_call(fn_name, fn_args or {}, catalog, uploaded_files)
            tool_events.append(result)
            prose = _clean_change_agent_text(text)
            messages.append({"role": "assistant", "content": prose or "(called tool)"})
            result_summary = json.dumps(
                {k: v for k, v in result.items() if k not in {"command", "artifact_previews", "cwd", "output_dir"}},
                ensure_ascii=False,
            )[:6000]
            messages.append({"role": "user", "content": f"Tool result for {fn_name}:\n{result_summary}\n\nNow answer the user's original question based on this data."})
        return {"text": "Reached tool-call limit.", "tool_events": tool_events}

    messages: list[dict[str, Any]] = [
        {"role": "system", "content": build_system_prompt(catalog, uploaded_files)}
    ]
    messages.extend(_normalize_history(history))
    messages.append({"role": "user", "content": user_message})

    tool_events: list[dict[str, Any]] = []
    tool_schemas = _tool_schemas(catalog)

    for _ in range(6):
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            tools=tool_schemas,
            tool_choice="auto",
            temperature=0.4,
        )
        message = response.choices[0].message
        message_dict = message.model_dump(exclude_none=True)
        messages.append(message_dict)

        if not message.tool_calls:
            inferred_events = tool_events or _infer_tool_details_from_text(message.content or "", catalog)
            return {
                "text": message.content or "No response returned.",
                "tool_events": inferred_events,
            }

        for tool_call in message.tool_calls:
            args = json.loads(tool_call.function.arguments or "{}")
            result = _dispatch_tool_call(
                name=tool_call.function.name,
                args=args,
                catalog=catalog,
                uploaded_files=uploaded_files,
            )
            tool_events.append(result)
            messages.append(
                {
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "name": tool_call.function.name,
                    "content": json.dumps(result, ensure_ascii=True),
                }
            )

    return {
        "text": "I hit the tool-calling limit before producing a final answer. Ask again with a narrower request.",
        "tool_events": tool_events,
    }


def _maybe_autorun_tool(
    user_message: str,
    history: list[dict[str, Any]],
    catalog: list[dict[str, Any]],
    uploaded_files: list[dict[str, Any]],
) -> dict[str, Any] | None:
    disclosure_args = _parse_disclosure_query(user_message)
    if not disclosure_args:
        disclosure_args = _parse_confirmation_from_history(user_message, history)

    if disclosure_args:
        result = _run_tool("influence_disclosure_tracker", disclosure_args, catalog, uploaded_files)
        return {
            "text": _compose_disclosure_response(result, disclosure_args),
            "tool_events": [result],
        }
    return None


def _tool_schemas(_catalog: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "type": "function",
            "function": {
                "name": "gnews_search",
                "description": "Search recent news articles via Google News. Use to gather current facts, monitor coverage, or research context before writing or advising.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "News search query"},
                        "period": {"type": "string", "description": "Time period: '1d', '7d', '30d', '1y' (default '7d')"},
                        "max_results": {"type": "integer", "description": "Max articles to return (default 8, max 20)"},
                    },
                    "required": ["query"],
                    "additionalProperties": False,
                },
            },
        },
        {
            "type": "function",
            "function": {
                "name": "brave_search",
                "description": "Search the open web via Brave Search. Use for research, fact-checking, finding primary sources, or any query where broader web coverage is needed.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Web search query"},
                        "count": {"type": "integer", "description": "Number of results (default 8, max 20)"},
                    },
                    "required": ["query"],
                    "additionalProperties": False,
                },
            },
        },
        {
            "type": "function",
            "function": {
                "name": "run_legislative_tracker",
                "description": "Run the Legislative Tracker to search bills, look up a specific bill by ID, or manage a watchlist.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string"},
                        "bill_id": {"type": "string"},
                        "state": {"type": "string"},
                        "year": {"type": "integer"},
                        "summarize": {"type": "boolean"},
                        "watchlist": {"type": "string", "enum": ["add", "remove", "list", "refresh"]},
                        "json": {"type": "boolean"},
                    },
                    "additionalProperties": False,
                },
            },
        },
        {
            "type": "function",
            "function": {
                "name": "run_influence_disclosure_tracker",
                "description": "Run the Influence Disclosure Tracker to look up LDA lobbying filings, FARA registrations, and IRS 990 data for any entity.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "entities": {
                            "oneOf": [
                                {"type": "string"},
                                {"type": "array", "items": {"type": "string"}},
                            ]
                        },
                        "from_date": {"type": "string"},
                        "to_date": {"type": "string"},
                        "filing_years": {
                            "oneOf": [
                                {"type": "string"},
                                {"type": "array", "items": {"type": "string"}},
                            ]
                        },
                        "filing_periods": {
                            "oneOf": [
                                {"type": "string"},
                                {"type": "array", "items": {"type": "string"}},
                            ]
                        },
                        "sources": {
                            "oneOf": [
                                {"type": "string"},
                                {"type": "array", "items": {"type": "string"}},
                            ]
                        },
                        "mode": {"type": "string", "enum": ["basic", "deep"]},
                        "search_field": {"type": "string", "enum": ["client", "registrant", "both"]},
                        "max_results": {"type": "integer"},
                    },
                    "required": ["entities"],
                    "additionalProperties": False,
                },
            },
        },
        {
            "type": "function",
            "function": {
                "name": "write_document",
                "description": "Package a fully written PA document as a downloadable DOCX + markdown artifact. Call this after you have written the complete document content. The content parameter should be the full document text in markdown.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string", "description": "Document title (used as filename and heading)"},
                        "document_type": {
                            "type": "string",
                            "enum": ["talking_points", "press_release", "op_ed", "speech", "memo", "social_media", "email", "pitch", "other"],
                            "description": "Type of PA document",
                        },
                        "content": {"type": "string", "description": "Full document content in markdown"},
                        "filename": {"type": "string", "description": "Optional custom base filename (no extension)"},
                    },
                    "required": ["title", "document_type", "content"],
                    "additionalProperties": False,
                },
            },
        },
    ]


def _dispatch_tool_call(
    name: str,
    args: dict[str, Any],
    catalog: list[dict[str, Any]],
    uploaded_files: list[dict[str, Any]],
) -> dict[str, Any]:
    if name == "gnews_search":
        return _run_gnews_search(args)
    if name == "brave_search":
        return _run_brave_search(args)
    if name == "write_document":
        return _run_write_document(args)
    if name in ("run_legislative_tracker", "legislative_tracker"):
        return _run_tool("legislative_tracker", args, catalog, uploaded_files)
    if name in ("run_influence_disclosure_tracker", "influence_disclosure_tracker"):
        # Normalize ChangeAgent param names (from/to → from_date/to_date)
        if "from" in args and "from_date" not in args:
            args["from_date"] = args.pop("from")
        if "to" in args and "to_date" not in args:
            args["to_date"] = args.pop("to")
        return _run_tool("influence_disclosure_tracker", args, catalog, uploaded_files)
    return {"ok": False, "error": f"Unknown tool call: {name}"}


def _get_tool_details(tool_id: str, catalog: list[dict[str, Any]]) -> dict[str, Any]:
    if tool_id == "media_clip_cleaner":
        return {
            "ok": True,
            "tool": {
                "id": "media_clip_cleaner",
                "name": "Media Clip Cleaner",
                "risk_level": "green",
                "frontend_path": "/media-clips",
                "required_inputs": ["raw_text or input_file_name"],
                "optional_inputs": ["mode", "llm_model", "fallback_local"],
                "when_to_use": TOOL_GUIDANCE["media_clip_cleaner"]["when_to_use"],
                "notes": TOOL_GUIDANCE["media_clip_cleaner"]["notes"],
            },
        }

    for tool in catalog:
        if tool["id"] == tool_id:
            return {
                "ok": True,
                "tool": {
                    "id": tool["id"],
                    "name": tool["name"],
                    "risk_level": tool.get("risk_level"),
                    "frontend_path": tool.get("frontend_path"),
                    "description": tool.get("description"),
                    "required_inputs": tool.get("inputs", {}).get("required", []),
                    "optional_inputs": tool.get("inputs", {}).get("optional", []),
                    "output_artifacts": tool.get("outputs", {}).get("artifacts", []),
                    "spec_summary": tool.get("spec_summary"),
                    "when_to_use": tool.get("assistant_guidance", {}).get("when_to_use"),
                    "notes": tool.get("assistant_guidance", {}).get("notes"),
                },
            }
    return {"ok": False, "error": f"Unknown tool_id: {tool_id}"}


def _run_tool(
    tool_id: str,
    arguments: dict[str, Any],
    catalog: list[dict[str, Any]],
    uploaded_files: list[dict[str, Any]],
) -> dict[str, Any]:
    try:
        if tool_id == "media_clips":
            return _run_media_clips(arguments)
        if tool_id == "media_clip_cleaner":
            return _run_media_clip_cleaner(arguments, uploaded_files)
        if tool_id == "influence_disclosure_tracker":
            return _run_disclosure_tracker(arguments)
        if tool_id == "hearing_memo_generator":
            return _run_hearing_memo(arguments, uploaded_files)
        if tool_id == "legislative_tracker":
            return _run_legislative_tracker(arguments)
        if tool_id == "messaging_matrix":
            return _run_messaging_matrix(arguments, uploaded_files)
        if tool_id == "media_list_builder":
            return _run_media_list_builder(arguments)
        if tool_id == "stakeholder_briefing":
            return _run_stakeholder_briefing(arguments, uploaded_files)
        if tool_id == "background_memo_generator":
            return _run_background_memo(arguments, uploaded_files)
        if tool_id == "stakeholder_map":
            return _run_stakeholder_map(arguments)
    except Exception as exc:
        return {"ok": False, "tool_id": tool_id, "error": str(exc)}

    return {"ok": False, "tool_id": tool_id, "error": f"Unsupported tool_id: {tool_id}"}


def _run_media_clips(arguments: dict[str, Any]) -> dict[str, Any]:
    topic = _pick(arguments, ["topic", "report_topic", "title"], required=True)
    queries = _pick_any(arguments, ["queries", "query", "search_queries", "query_list"])
    if isinstance(queries, list):
        query_value = ",".join(q.strip() for q in queries if q and q.strip())
    else:
        query_value = str(queries or "").strip()
    if not query_value:
        raise ValueError("media_clips requires `queries` as a string or list of query strings.")

    outdir = _output_dir("media_clips", topic)
    script = TOOLKIT_ROOT / "tools" / "media_clips" / "execution" / "generate_clips.py"
    cmd = [
        sys.executable,
        str(script),
        "--topic",
        topic,
        "--queries",
        query_value,
        "--output-dir",
        str(outdir),
        "--no-email",
    ]
    _optional_arg(cmd, "--period", arguments.get("period"))
    _optional_arg(cmd, "--suffix", arguments.get("suffix"))
    _optional_arg(cmd, "--since", arguments.get("since"))
    _optional_arg(cmd, "--target-date", arguments.get("target_date"))

    if arguments.get("all_sources"):
        cmd.append("--all-sources")
    custom_sources = arguments.get("custom_sources")
    if custom_sources:
        if isinstance(custom_sources, list):
            custom_sources = ",".join(custom_sources)
        _optional_arg(cmd, "--custom-sources", custom_sources)

    return _execute_cli("media_clips", cmd, outdir, cwd=script.parent)


def _run_media_clip_cleaner(arguments: dict[str, Any], uploaded_files: list[dict[str, Any]]) -> dict[str, Any]:
    raw_text = _pick(arguments, ["raw_text", "article_text", "text", "content"]) or ""
    input_path = None
    if not raw_text:
        input_path = _resolve_uploaded_file(
            arguments,
            uploaded_files,
            ["input_file_name", "input_path", "file_name", "source_file"],
        )
        if not input_path:
            raise ValueError("media_clip_cleaner requires `raw_text` or `input_file_name`.")

    outdir = _output_dir("media_clip_cleaner", "cleaned_clip")
    outfile = outdir / "cleaned_clip.md"
    script = TOOLKIT_ROOT / "tools" / "media_clip_cleaner" / "execution" / "clean_clip.py"
    cmd = [
        sys.executable,
        str(script),
        "--output-file",
        str(outfile),
        "--mode",
        str(arguments.get("mode") or "local"),
    ]
    _optional_arg(cmd, "--title", arguments.get("title"))
    if raw_text:
        cmd.extend(["--raw-text", raw_text])
    else:
        cmd.extend(["--input-file", str(input_path)])
    _optional_arg(cmd, "--llm-model", arguments.get("llm_model"))
    if arguments.get("fallback_local"):
        cmd.append("--fallback-local")

    return _execute_cli("media_clip_cleaner", cmd, outdir, cwd=script.parent)


def _run_disclosure_tracker(arguments: dict[str, Any]) -> dict[str, Any]:
    entities = _pick_any(arguments, ["entities", "entity", "entity_names"])
    if isinstance(entities, list):
        entities_value = ",".join(str(entity).strip() for entity in entities if str(entity).strip())
    else:
        entities_value = str(entities or "").strip()
    if not entities_value:
        raise ValueError("influence_disclosure_tracker requires `entities`.")

    outdir = _output_dir("influence_disclosure_tracker", entities_value)
    script = TOOLKIT_ROOT / "tools" / "influence_disclosure_tracker" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--entities",
        entities_value,
        "--out",
        str(outdir),
        "--format",
        str(arguments.get("format") or "csv,md"),
    ]
    _optional_arg(cmd, "--from", arguments.get("from_date"))
    _optional_arg(cmd, "--to", arguments.get("to_date"))
    _optional_arg(cmd, "--filing-years", _csv_value(arguments.get("filing_years")))
    _optional_arg(cmd, "--filing-periods", _csv_value(arguments.get("filing_periods")))
    _optional_arg(cmd, "--sources", _csv_value(arguments.get("sources")))
    _optional_arg(cmd, "--mode", arguments.get("mode"))
    _optional_arg(cmd, "--search-field", arguments.get("search_field"))
    _optional_arg(cmd, "--max-results", _stringify(arguments.get("max_results")))
    if arguments.get("debug"):
        cmd.append("--debug")

    return _execute_cli("influence_disclosure_tracker", cmd, outdir, cwd=script.parent, timeout=240)


def _run_hearing_memo(arguments: dict[str, Any], uploaded_files: list[dict[str, Any]]) -> dict[str, Any]:
    input_path = _resolve_uploaded_file(
        arguments,
        uploaded_files,
        ["input_file_name", "input_path", "file_name", "transcript_file", "transcript_file_name"],
        required=True,
    )

    outdir = _output_dir("hearing_memo_generator", Path(input_path).stem)
    docx_path = outdir / "hearing_memo.docx"
    text_path = outdir / "hearing_memo.md"
    json_path = outdir / "hearing_memo_bundle.json"
    script = TOOLKIT_ROOT / "tools" / "hearing_memo" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--input",
        str(input_path),
        "--output",
        str(docx_path),
        "--text-output",
        str(text_path),
        "--json-output",
        str(json_path),
    ]
    _optional_arg(cmd, "--from", arguments.get("memo_from"))
    _optional_arg(cmd, "--memo-date", arguments.get("memo_date"))
    _optional_arg(cmd, "--subject", arguments.get("subject"))
    _optional_arg(cmd, "--confidentiality-footer", arguments.get("confidentiality_footer"))
    _optional_arg(cmd, "--hearing-title", arguments.get("hearing_title"))
    _optional_arg(cmd, "--hearing-date", arguments.get("hearing_date"))
    _optional_arg(cmd, "--hearing-time", arguments.get("hearing_time"))
    _optional_arg(cmd, "--committee", arguments.get("committee"))
    if arguments.get("verbose"):
        cmd.append("--verbose")

    return _execute_cli("hearing_memo_generator", cmd, outdir, cwd=script.parent, timeout=300)


def _run_legislative_tracker(arguments: dict[str, Any]) -> dict[str, Any]:
    query = _pick(arguments, ["query", "search_query", "topic"])
    bill_id = _pick(arguments, ["bill_id", "bill"])
    watchlist = _pick(arguments, ["watchlist", "watchlist_action"])
    if not query and not bill_id and not watchlist:
        raise ValueError("legislative_tracker requires `query`, `bill_id`, or `watchlist`.")

    slug_source = bill_id or query or watchlist
    outdir = _output_dir("legislative_tracker", slug_source)
    script = TOOLKIT_ROOT / "tools" / "legislative_tracker" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--out",
        str(outdir),
    ]
    _optional_arg(cmd, "--query", query)
    _optional_arg(cmd, "--state", _pick(arguments, ["state", "jurisdiction"]))
    _optional_arg(cmd, "--year", _stringify(arguments.get("year")))
    _optional_arg(cmd, "--bill-id", bill_id)
    _optional_arg(cmd, "--watchlist", watchlist)
    if arguments.get("summarize"):
        cmd.append("--summarize")
    if arguments.get("json"):
        cmd.append("--json")

    return _execute_cli("legislative_tracker", cmd, outdir, cwd=script.parent, timeout=240)


def _run_messaging_matrix(arguments: dict[str, Any], uploaded_files: list[dict[str, Any]]) -> dict[str, Any]:
    position = _pick(arguments, ["position", "policy_position", "core_position"], required=True)
    outdir = _output_dir("messaging_matrix", position)
    script = TOOLKIT_ROOT / "tools" / "messaging_matrix" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--position",
        position,
        "--out",
        str(outdir),
        "--json",
    ]
    _optional_arg(cmd, "--context", _pick(arguments, ["context", "background", "notes"]))
    _optional_arg(cmd, "--organization", arguments.get("organization"))
    _optional_arg(cmd, "--audience", _pick(arguments, ["audience", "target_audience"]))
    _optional_arg(cmd, "--core-messages", _pick(arguments, ["core_messages", "messages"]))
    _optional_arg(cmd, "--facts", arguments.get("facts"))
    _optional_arg(cmd, "--variants", _csv_value(arguments.get("variants")))
    context_file = _resolve_uploaded_file(
        arguments,
        uploaded_files,
        ["context_file_name", "context_file", "file_name", "source_file"],
    )
    if context_file:
        cmd.extend(["--context-file", str(context_file)])

    return _execute_cli("messaging_matrix", cmd, outdir, cwd=script.parent, timeout=240)


def _run_media_list_builder(arguments: dict[str, Any]) -> dict[str, Any]:
    issue = _pick(arguments, ["issue", "topic", "policy_issue"], required=True)
    outdir = _output_dir("media_list_builder", issue)
    script = TOOLKIT_ROOT / "tools" / "media_list_builder" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--issue",
        issue,
        "--out",
        str(outdir),
        "--json",
    ]
    _optional_arg(cmd, "--location", arguments.get("location"))
    _optional_arg(cmd, "--media-types", _csv_value(arguments.get("media_types")))
    _optional_arg(cmd, "--num-contacts", _stringify(arguments.get("num_contacts")))

    return _execute_cli("media_list_builder", cmd, outdir, cwd=script.parent, timeout=300)


def _run_stakeholder_briefing(arguments: dict[str, Any], uploaded_files: list[dict[str, Any]]) -> dict[str, Any]:
    name = _pick(arguments, ["name", "stakeholder_name", "person", "stakeholder"], required=True)
    purpose = _pick(arguments, ["purpose", "meeting_purpose", "objective"], required=True)

    outdir = _output_dir("stakeholder_briefing", name)
    script = TOOLKIT_ROOT / "tools" / "stakeholder_briefing" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--name",
        str(name),
        "--purpose",
        str(purpose),
        "--out",
        str(outdir),
        "--json",
    ]
    _optional_arg(cmd, "--organization", arguments.get("organization"))
    _optional_arg(cmd, "--your-org", arguments.get("your_org"))
    _optional_arg(cmd, "--context", _pick(arguments, ["context", "background", "notes"]))
    context_file = _resolve_uploaded_file(
        arguments,
        uploaded_files,
        ["context_file_name", "context_file", "file_name", "source_file"],
    )
    if context_file:
        cmd.extend(["--context-file", str(context_file)])
    if arguments.get("include_disclosures") is False:
        cmd.append("--no-disclosures")
    if arguments.get("include_news") is False:
        cmd.append("--no-news")

    return _execute_cli("stakeholder_briefing", cmd, outdir, cwd=script.parent, timeout=300)


def _run_background_memo(arguments: dict[str, Any], uploaded_files: list[dict[str, Any]]) -> dict[str, Any]:
    subject = _pick(arguments, ["subject", "topic", "entity"], required=True)
    sections = _pick_any(arguments, ["sections", "headings", "section_headings"])
    if isinstance(sections, str):
        section_values = [line.strip() for line in sections.splitlines() if line.strip()]
    else:
        section_values = [str(section).strip() for section in (sections or []) if str(section).strip()]
    if not section_values:
        raise ValueError("background_memo_generator requires `sections` as a list or newline-delimited string.")

    context = _pick(arguments, ["context", "background", "notes"]) or ""
    context_file = _resolve_uploaded_file(
        arguments,
        uploaded_files,
        ["context_file_name", "context_file", "file_name", "source_file"],
    )
    if context_file:
        file_text = _read_text_preview(Path(context_file), limit=12000)
        context = "\n\n".join(part for part in [context, f"Source file excerpt:\n{file_text}"] if part).strip()

    outdir = _output_dir("background_memo_generator", subject)
    docx_path = outdir / "background_memo.docx"
    md_path = outdir / "background_memo.md"
    script = TOOLKIT_ROOT / "tools" / "background_memo" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--subject",
        subject,
        "--out",
        str(docx_path),
        "--md-out",
        str(md_path),
        "--sections",
        *section_values,
    ]
    _optional_arg(cmd, "--context", context)
    _optional_arg(cmd, "--date", arguments.get("date"))

    return _execute_cli("background_memo_generator", cmd, outdir, cwd=script.parent, timeout=300)


def _run_stakeholder_map(arguments: dict[str, Any]) -> dict[str, Any]:
    policy_issue = _pick(arguments, ["policy_issue", "issue", "topic"], required=True)
    outdir = _output_dir("stakeholder_map", policy_issue)
    script = TOOLKIT_ROOT / "tools" / "stakeholder_map" / "execution" / "run.py"
    cmd = [
        sys.executable,
        str(script),
        "--policy_issue",
        policy_issue,
        "--out",
        str(outdir),
    ]
    _optional_arg(cmd, "--scope", arguments.get("scope"))
    _optional_arg(cmd, "--state", arguments.get("state"))
    include_types = arguments.get("include_types")
    if include_types:
        cmd.extend(["--include_types", *_list_value(include_types)])
    if arguments.get("no_graph"):
        cmd.append("--no_graph")

    return _execute_cli("stakeholder_map", cmd, outdir, cwd=script.parent, timeout=300)


def _run_gnews_search(arguments: dict[str, Any]) -> dict[str, Any]:
    query = arguments.get("query", "").strip()
    if not query:
        return {"ok": False, "tool_id": "gnews_search", "error": "query is required"}
    period = arguments.get("period") or "7d"
    max_results = min(int(arguments.get("max_results") or 8), 20)
    try:
        from gnews import GNews
        gn = GNews(language="en", country="US", period=period, max_results=max_results)
        articles = gn.get_news(query) or []
        results = [
            {
                "title": a.get("title", ""),
                "description": a.get("description", ""),
                "url": a.get("url", ""),
                "published_date": a.get("published date", ""),
                "publisher": (a.get("publisher") or {}).get("title", ""),
            }
            for a in articles
        ]
        return {"ok": True, "tool_id": "gnews_search", "query": query, "results": results}
    except Exception as exc:
        return {"ok": False, "tool_id": "gnews_search", "error": str(exc)}


def _run_brave_search(arguments: dict[str, Any]) -> dict[str, Any]:
    query = arguments.get("query", "").strip()
    if not query:
        return {"ok": False, "tool_id": "brave_search", "error": "query is required"}
    count = min(int(arguments.get("count") or 8), 20)
    api_key = os.environ.get("BRAVE_SEARCH_API_KEY", "").strip()
    if not api_key:
        return {"ok": False, "tool_id": "brave_search", "error": "BRAVE_SEARCH_API_KEY is not set"}
    try:
        resp = requests.get(
            "https://api.search.brave.com/res/v1/web/search",
            headers={"Accept": "application/json", "X-Subscription-Token": api_key},
            params={"q": query, "count": count},
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        web_results = data.get("web", {}).get("results", [])
        results = [
            {
                "title": r.get("title", ""),
                "description": r.get("description", ""),
                "url": r.get("url", ""),
            }
            for r in web_results
        ]
        return {"ok": True, "tool_id": "brave_search", "query": query, "results": results}
    except Exception as exc:
        return {"ok": False, "tool_id": "brave_search", "error": str(exc)}


def _run_write_document(arguments: dict[str, Any]) -> dict[str, Any]:
    title = arguments.get("title", "Document").strip()
    content = arguments.get("content", "").strip()
    doc_type = arguments.get("document_type", "other")
    filename_base = _slugify(arguments.get("filename") or title) or "document"

    outdir = _output_dir("remy_writer", title)
    md_path = outdir / f"{filename_base}.md"
    docx_path = outdir / f"{filename_base}.docx"

    # Write markdown
    full_md = f"# {title}\n\n{content}"
    md_path.write_text(full_md, encoding="utf-8")

    # Write DOCX
    try:
        _markdown_to_docx(title, content, docx_path)
    except Exception as exc:
        return {"ok": False, "tool_id": "write_document", "error": f"DOCX generation failed: {exc}"}

    artifacts = _collect_artifacts(outdir)
    return {
        "ok": True,
        "tool_id": "write_document",
        "document_type": doc_type,
        "title": title,
        "output_dir": str(outdir),
        "artifacts": artifacts,
        "artifact_previews": _artifact_previews(artifacts),
        "frontend_path": None,
    }


def _markdown_to_docx(title: str, content: str, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ", "+ ")):
            para = doc.add_paragraph(style="List Bullet")
            _apply_inline_markdown(para, line[2:].strip())
        elif line.startswith(("  - ", "  * ")):
            para = doc.add_paragraph(style="List Bullet 2")
            _apply_inline_markdown(para, line[4:].strip())
        elif line == "":
            doc.add_paragraph("")
        elif line.startswith("---") or line.startswith("***"):
            doc.add_paragraph("─" * 40)
        else:
            para = doc.add_paragraph()
            _apply_inline_markdown(para, line)
            para.runs[0].font.size = Pt(11) if para.runs else None

    doc.save(str(output_path))


def _apply_inline_markdown(para, text: str) -> None:
    """Parse **bold** and *italic* inline markdown and add styled runs."""
    import re as _re
    pattern = _re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif full.startswith("*"):
            run = para.add_run(m.group(3))
            run.italic = True
        elif full.startswith("`"):
            run = para.add_run(m.group(4))
            run.font.name = "Courier New"
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])
    if not para.runs:
        para.add_run(text)


def _execute_cli(
    tool_id: str,
    cmd: list[str],
    outdir: Path,
    cwd: Path,
    timeout: int = 180,
) -> dict[str, Any]:
    outdir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    result = subprocess.run(
        cmd,
        cwd=str(cwd),
        capture_output=True,
        text=True,
        timeout=timeout,
        env=env,
    )

    artifacts = _collect_artifacts(outdir)
    payload = {
        "ok": result.returncode == 0,
        "tool_id": tool_id,
        "command": cmd,
        "cwd": str(cwd),
        "output_dir": str(outdir),
        "returncode": result.returncode,
        "stdout": _truncate(result.stdout),
        "stderr": _truncate(result.stderr),
        "artifacts": artifacts,
        "artifact_previews": _artifact_previews(artifacts),
        "frontend_path": _FRONTEND_PATHS.get(tool_id),
    }
    if not payload["ok"] and not payload["stderr"]:
        payload["error"] = "Tool exited with a non-zero status."
    return payload


def _collect_artifacts(outdir: Path) -> list[dict[str, Any]]:
    if not outdir.exists():
        return []

    artifacts = []
    for path in sorted(p for p in outdir.rglob("*") if p.is_file()):
        artifacts.append(
            {
                "name": path.name,
                "path": str(path),
                "suffix": path.suffix.lower(),
                "size_bytes": path.stat().st_size,
            }
        )
    return artifacts[:20]


def _artifact_previews(artifacts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    previews = []
    for artifact in artifacts:
        suffix = artifact["suffix"]
        if suffix not in {".md", ".txt", ".json", ".csv"}:
            continue
        path = Path(artifact["path"])
        previews.append(
            {
                "name": artifact["name"],
                "path": artifact["path"],
                "preview": _read_text_preview(path, limit=2500),
            }
        )
        if len(previews) >= 3:
            break
    return previews


def _read_spec_summary(spec_path: Path) -> str:
    if not spec_path.exists():
        return ""

    lines = []
    for raw_line in spec_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("#") and lines:
            break
        lines.append(line)
        if len(" ".join(lines)) > 450:
            break
    return " ".join(lines)[:450]


def _infer_tool_details_from_text(text: str, catalog: list[dict[str, Any]]) -> list[dict[str, Any]]:
    lowered = text.lower()
    matches = []
    for tool in catalog:
        if tool["name"].lower() in lowered:
            matches.append(_get_tool_details(tool["id"], catalog))
        elif tool["id"].replace("_", " ") in lowered:
            matches.append(_get_tool_details(tool["id"], catalog))
    return matches[:2]


def _parse_disclosure_query(user_message: str) -> dict[str, Any] | None:
    text = " ".join(user_message.strip().split())
    lowered = text.lower()

    if "lobby" not in lowered:
        return None

    entity = None
    patterns = [
        r"(?:can you please\s+|please\s+|hi remy,?\s+|remy,?\s+|look\s+|find out\s+|tell me\s+)?how much\s+(?P<entity>.+?)\s+spent\b",
        r"what did\s+(?P<entity>.+?)\s+spend\b",
        r"how much lobbying did\s+(?P<entity>.+?)\s+do",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            entity = match.group("entity")
            break

    period = _parse_quarter_range(text)
    if not entity or not period:
        return None

    return {
        "entities": _clean_entity_name(entity),
        "from_date": period["from_date"],
        "to_date": period["to_date"],
        "sources": ["lda"],
        "mode": "basic",
    }


def _parse_confirmation_from_history(
    user_message: str,
    history: list[dict[str, Any]],
) -> dict[str, Any] | None:
    lowered = user_message.strip().lower()
    if lowered not in {
        "yes",
        "yes.",
        "yes, proceed",
        "yes, proceed.",
        "yes, proceed with those inputs",
        "yes, proceed with those inputs.",
        "proceed",
        "proceed.",
        "go ahead",
        "go ahead.",
    }:
        return None

    for item in reversed(history):
        if item.get("role") != "assistant":
            continue
        content = item.get("content") or ""
        entity_match = re.search(r'Entities:\s*"([^"]+)"', content, flags=re.IGNORECASE)
        from_match = re.search(r'From date:\s*"(\d{4}-\d{2}-\d{2})"', content, flags=re.IGNORECASE)
        to_match = re.search(r'To date:\s*"(\d{4}-\d{2}-\d{2})"', content, flags=re.IGNORECASE)
        if entity_match and from_match and to_match:
            return {
                "entities": entity_match.group(1).strip(),
                "from_date": from_match.group(1),
                "to_date": to_match.group(1),
                "sources": ["lda"],
                "mode": "basic",
            }
    return None


def _parse_quarter_range(text: str) -> dict[str, str] | None:
    quarter_match = re.search(r"\bq([1-4])\s*(20\d{2})\b", text, flags=re.IGNORECASE)
    if quarter_match:
        quarter = int(quarter_match.group(1))
        year = int(quarter_match.group(2))
        return _quarter_to_dates(year, quarter)

    named_match = re.search(
        r"\b(first|1st|second|2nd|third|3rd|fourth|4th|four|last)\s+quarter\s+of\s+(20\d{2})\b",
        text,
        flags=re.IGNORECASE,
    )
    if named_match:
        quarter_token = named_match.group(1).lower()
        year = int(named_match.group(2))
        quarter_map = {
            "first": 1,
            "1st": 1,
            "second": 2,
            "2nd": 2,
            "third": 3,
            "3rd": 3,
            "fourth": 4,
            "4th": 4,
            "four": 4,
            "last": 4,
        }
        return _quarter_to_dates(year, quarter_map[quarter_token])

    return None


def _quarter_to_dates(year: int, quarter: int) -> dict[str, str]:
    starts = {
        1: (f"{year}-01-01", f"{year}-03-31"),
        2: (f"{year}-04-01", f"{year}-06-30"),
        3: (f"{year}-07-01", f"{year}-09-30"),
        4: (f"{year}-10-01", f"{year}-12-31"),
    }
    from_date, to_date = starts[quarter]
    return {"from_date": from_date, "to_date": to_date}


def _clean_entity_name(value: str) -> str:
    cleaned = re.sub(r"^(the)\s+", "", value.strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+in\s+the\s+.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+on\s+lobbying.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"[?.!,]+$", "", cleaned)
    if cleaned.islower():
        return cleaned.strip().title()
    return cleaned.strip()


def _compose_disclosure_response(result: dict[str, Any], arguments: dict[str, Any]) -> str:
    entity = arguments["entities"]
    from_date = arguments["from_date"]
    to_date = arguments["to_date"]

    if not result.get("ok"):
        error = result.get("error") or result.get("stderr") or "unknown error"
        return (
            f"I tried to run the Influence Disclosure Tracker for {entity} from {from_date} through {to_date}, "
            f"but it failed: {error}"
        )

    summary = _summarize_lda_filings(
        result,
        entity=entity,
        filing_period=_filing_period_from_dates(from_date, to_date),
    )
    if not summary["rows"]:
        return (
            f"I ran the Influence Disclosure Tracker for {entity} from {from_date} through {to_date}, "
            "but I did not find billable LDA filing rows with reported amounts in that range."
        )

    direct_lines = "\n".join(
        f"- {name}: ${amount:,.0f}"
        for name, amount in summary["direct_payers"]
    )
    response = (
        f"I ran the Influence Disclosure Tracker for {entity} from {from_date} through {to_date}. "
        f"For that quarter, the direct reported lobbying total tied to {entity} was ${summary['direct_total']:,.0f}.\n\n"
        f"Who appears to have been paid directly by {entity} in the LDA filings:\n{direct_lines}"
    )

    if summary["indirect_payers"]:
        indirect_lines = "\n".join(
            f"- {name}: ${amount:,.0f}"
            for name, amount in summary["indirect_payers"]
        )
        response += (
            "\n\nThere is also an indirect/subcontractor filing that appears to sit on top of another firm's engagement:\n"
            f"{indirect_lines}\n\n"
            "I am not adding that indirect amount into the direct client-spend total to avoid double counting."
        )

    response += (
        "\n\nThis is based on the underlying LDA filing data returned by the tool. "
        "Treat it as medium-risk output and review the filings before external use. "
        "[Search LDA filings →](https://lda.senate.gov/filings/public/filing/search/)"
    )
    return response


def _summarize_lda_filings(
    result: dict[str, Any],
    entity: str,
    filing_period: str | None,
) -> dict[str, Any]:
    lda_path = None
    for artifact in result.get("artifacts", []):
        if artifact["name"] == "lda_filings.csv":
            lda_path = Path(artifact["path"])
            break

    if not lda_path or not lda_path.exists():
        return {
            "direct_total": 0.0,
            "direct_payers": [],
            "indirect_payers": [],
            "rows": [],
        }

    rows = []
    direct_totals: dict[str, float] = {}
    indirect_totals: dict[str, float] = {}
    entity_norm = _normalize_entity(entity)
    with open(lda_path, encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row_period = (row.get("filing_period") or "").strip().lower()
            if filing_period and row_period != filing_period:
                continue
            amount_raw = (row.get("amount") or "").strip()
            if not amount_raw:
                continue
            try:
                amount = float(amount_raw)
            except ValueError:
                continue
            if amount <= 0:
                continue
            rows.append(row)
            registrant = (row.get("registrant_name") or "Unknown").strip()
            client_name = (row.get("client_name") or "").strip()
            client_norm = _normalize_entity(client_name)
            if entity_norm and entity_norm == client_norm:
                direct_totals[registrant] = direct_totals.get(registrant, 0.0) + amount
            else:
                indirect_totals[registrant] = indirect_totals.get(registrant, 0.0) + amount

    direct_payers = sorted(direct_totals.items(), key=lambda item: item[1], reverse=True)
    indirect_payers = sorted(indirect_totals.items(), key=lambda item: item[1], reverse=True)
    direct_total = sum(amount for _, amount in direct_payers)
    return {
        "direct_total": direct_total,
        "direct_payers": direct_payers,
        "indirect_payers": indirect_payers,
        "rows": rows,
    }


def _filing_period_from_dates(from_date: str, to_date: str) -> str | None:
    if len(from_date) != 10 or len(to_date) != 10:
        return None
    year = from_date[:4]
    quarter_map = {
        (f"{year}-01-01", f"{year}-03-31"): "first_quarter",
        (f"{year}-04-01", f"{year}-06-30"): "second_quarter",
        (f"{year}-07-01", f"{year}-09-30"): "third_quarter",
        (f"{year}-10-01", f"{year}-12-31"): "fourth_quarter",
    }
    return quarter_map.get((from_date, to_date))


def _normalize_entity(value: str) -> str:
    normalized = value.upper()
    normalized = normalized.replace(",", " ")
    normalized = normalized.replace(".", " ")
    normalized = normalized.replace("LLC", " ")
    normalized = normalized.replace("INC", " ")
    normalized = normalized.replace("CORP", " ")
    normalized = normalized.replace("CORPORATION", " ")
    normalized = normalized.replace("COMPANY", " ")
    normalized = normalized.replace("PBC", " ")
    normalized = normalized.replace("ON BEHALF OF", " ")
    normalized = " ".join(normalized.split())
    return normalized.strip()


def _normalize_history(history: list[dict[str, Any]]) -> list[dict[str, Any]]:
    normalized = []
    for item in history:
        role = item.get("role")
        content = item.get("content")
        if role in {"user", "assistant"} and isinstance(content, str):
            normalized.append({"role": role, "content": content})
    return normalized


def _resolve_uploaded_file(
    arguments: dict[str, Any],
    uploaded_files: list[dict[str, Any]],
    keys: list[str],
    required: bool = False,
) -> str | None:
    file_lookup = {
        item["name"]: item["path"]
        for item in uploaded_files
    }
    file_lookup.update(
        {
            item["path"]: item["path"]
            for item in uploaded_files
        }
    )

    for key in keys:
        value = arguments.get(key)
        if not value:
            continue
        resolved = file_lookup.get(str(value))
        if resolved:
            return resolved
        if Path(str(value)).exists():
            return str(Path(str(value)).resolve())

    if required:
        names = ", ".join(sorted(file_lookup.keys())[:20]) or "none"
        raise ValueError(f"Required uploaded file was not found. Available file references: {names}")
    return None


def _output_dir(tool_id: str, slug_source: str) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = _slugify(slug_source)[:48] or tool_id
    outdir = OUTPUT_ROOT / tool_id / f"{timestamp}_{slug}"
    outdir.mkdir(parents=True, exist_ok=True)
    return outdir


def _slugify(value: str) -> str:
    chars = []
    for char in str(value).lower():
        if char.isalnum():
            chars.append(char)
        elif char in {" ", "-", "_", "."}:
            chars.append("_")
    slug = "".join(chars)
    while "__" in slug:
        slug = slug.replace("__", "_")
    return slug.strip("_")


def _truncate(text: str | None, limit: int = 4000) -> str:
    if not text:
        return ""
    if len(text) <= limit:
        return text
    return text[:limit] + "\n...[truncated]"


def _read_text_preview(path: Path, limit: int = 2500) -> str:
    try:
        if path.suffix.lower() == ".json":
            data = json.loads(path.read_text(encoding="utf-8"))
            text = json.dumps(data, indent=2)[:limit]
            return text
        return path.read_text(encoding="utf-8", errors="replace")[:limit]
    except Exception as exc:
        return f"[Preview unavailable: {exc}]"


def _optional_arg(cmd: list[str], flag: str, value: Any) -> None:
    if value is None:
        return
    value_str = str(value).strip()
    if not value_str:
        return
    cmd.extend([flag, value_str])


def _require(arguments: dict[str, Any], key: str) -> str:
    value = arguments.get(key)
    if value is None or not str(value).strip():
        raise ValueError(f"Missing required argument: {key}")
    return str(value).strip()


def _pick(arguments: dict[str, Any], keys: list[str], required: bool = False) -> str | None:
    for key in keys:
        value = arguments.get(key)
        if value is None:
            continue
        value_str = str(value).strip()
        if value_str:
            return value_str
    if required:
        raise ValueError(f"Missing required argument. Expected one of: {', '.join(keys)}")
    return None


def _pick_any(arguments: dict[str, Any], keys: list[str]) -> Any:
    for key in keys:
        value = arguments.get(key)
        if value is None:
            continue
        if isinstance(value, list) and value:
            return value
        if isinstance(value, str) and value.strip():
            return value
        if not isinstance(value, str):
            return value
    return None


def _csv_value(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, list):
        cleaned = [str(item).strip() for item in value if str(item).strip()]
        return ",".join(cleaned) if cleaned else None
    value_str = str(value).strip()
    return value_str or None


def _list_value(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    value_str = str(value).strip()
    return [value_str] if value_str else []


def _stringify(value: Any) -> str | None:
    if value is None or value == "":
        return None
    return str(value)
