"""
Background Memo DOCX Exporter
================================
Produces a DOCX matching the exact house style observed in the example memos:

Structure:
  DATE: <date>          ← Normal style, plain
  SUBJECT: <subject>    ← Normal style, plain
  [blank line]
  Overview              ← Normal style, plain (section label)
  <overview paragraph>  ← Normal style, plain
  [blank line]
  Fast Facts            ← Normal style, plain (section label)
  • Bold sentence.      ← List Paragraph style, bold text
  • Bold sentence.
  [blank line]
  <Section Heading>     ← Normal style, plain
  [Sub-heading if any]  ← Normal style, bold
  <paragraph>           ← Normal style, plain
  ...
  Relevant Links        ← Normal style, plain
  • Label — URL         ← List Paragraph style, plain
"""

import sys
from datetime import date
from pathlib import Path

from docx import Document

# Schema is in the same directory; ensure it's importable both via CLI and FastAPI.
sys.path.insert(0, str(Path(__file__).parent))
from schema import BackgroundMemoResult  # noqa: E402


def _add_normal(doc: Document, text: str, bold: bool = False) -> None:
    """Add a Normal-style paragraph."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = bold


def _add_list_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """Add a List Paragraph style bullet."""
    p = doc.add_paragraph(style="List Paragraph")
    run = p.add_run(text)
    run.bold = bold


def _add_blank(doc: Document) -> None:
    doc.add_paragraph(style="Normal")


def export_docx(result: "BackgroundMemoResult | dict", output_path: str, memo_date: str = "") -> None:
    """
    Write a background memo DOCX to output_path.

    Accepts either a BackgroundMemoResult model or a plain dict (e.g. from model_dump()).
    Normalizes to BackgroundMemoResult so field access is always typed and explicit.

    Args:
        result:      Output from generator.generate_memo() or a BackgroundMemoResult instance.
        output_path: Full path to write the .docx file.
        memo_date:   Date string (default: today formatted as "Month DD, YYYY").
    """
    if isinstance(result, dict):
        result = BackgroundMemoResult(**result)

    doc = Document()

    d = memo_date or date.today().strftime("%B %d, %Y")

    # ── Header block ─────────────────────────────────────────────────────────
    _add_normal(doc, f"DATE:\t\t{d}")
    _add_normal(doc, f"SUBJECT:\t{result.subject} Background Memo")
    _add_blank(doc)

    # ── Overview ─────────────────────────────────────────────────────────────
    _add_normal(doc, "Overview")
    _add_normal(doc, result.overview)
    _add_blank(doc)

    # ── Fast Facts ───────────────────────────────────────────────────────────
    _add_normal(doc, "Fast Facts")
    for fact in result.fast_facts:
        _add_list_paragraph(doc, fact, bold=True)
    _add_blank(doc)

    # ── Content sections ─────────────────────────────────────────────────────
    for section in result.sections:
        _add_normal(doc, section.heading)
        for sub in section.subsections:
            if sub.heading:
                _add_normal(doc, sub.heading, bold=True)
            for para in sub.paragraphs:
                _add_normal(doc, para)
        _add_blank(doc)

    # ── Links ─────────────────────────────────────────────────────────────────
    _add_normal(doc, "Links")
    for link in result.links:
        _add_list_paragraph(doc, f"{link.label} — {link.url}")

    doc.save(output_path)
