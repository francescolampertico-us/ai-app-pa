"""
Shared DOCX brand module — Strategitect.
Direction: Obsidian.

All values read from brand.toml. Fallbacks hardcoded to Obsidian defaults
so the module works without brand.toml present.

Usage in a tool export script:
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).resolve().parents[2]))  # → tools/shared/
    from brand.docx_brand import new_branded_doc, apply_cover, section_heading, ...
"""
from __future__ import annotations
from pathlib import Path

# ─── Locate brand.toml ────────────────────────────────────────────────────────
_HERE         = Path(__file__).resolve().parent
_PROJECT_ROOT = _HERE.parents[3]          # brand/ → shared/ → tools/ → toolkit/ → root
_BRAND_TOML   = _PROJECT_ROOT / "brand-toolkit" / "brand.toml"

# ─── Load TOML ────────────────────────────────────────────────────────────────
_cfg: dict = {}
if _BRAND_TOML.exists():
    try:
        import tomllib                    # Python 3.11+
    except ImportError:
        try:
            import tomli as tomllib       # type: ignore
        except ImportError:
            tomllib = None                # type: ignore
    if tomllib is not None:
        try:
            with open(_BRAND_TOML, "rb") as f:
                _cfg = tomllib.load(f)
        except Exception:
            pass


def _c(section: str, key: str, default: str) -> str:
    parts = section.split(".")
    d: dict | str = _cfg
    for p in parts:
        d = d.get(p, {}) if isinstance(d, dict) else {}
    return d.get(key, default) if isinstance(d, dict) else default   # type: ignore


# ─── Brand identity ───────────────────────────────────────────────────────────
COMPANY_NAME = _c("company", "name",        "Strategitect")
FOOTER_TEXT  = _c("company", "footer_text", "Strategitect — Confidential")
COVER_BADGE  = _c("company", "cover_badge", "CONFIDENTIAL")
SLOGAN       = _c("company", "slogan",      "Architecture for Public Affairs Strategy")

# ─── Fonts (python-docx uses system font names) ───────────────────────────────
_docx_fonts = _cfg.get("fonts", {}).get("docx", {})
FONT_DISPLAY = _docx_fonts.get("display", "DM Serif Display")
FONT_BODY    = _docx_fonts.get("body",    "Inter")
FONT_MONO    = _docx_fonts.get("mono",    "Courier New")

# ─── Color tokens (Obsidian) ──────────────────────────────────────────────────
_col  = _cfg.get("colors", {})
_zinc = _col.get("zinc", {})
_comp = _cfg.get("components", {})
_call = _comp.get("callout", {})
_tbl  = _comp.get("table", {})

_HEX_PRIMARY = _col.get("primary",  "6D28D9")   # Violet-700
_HEX_ACCENT  = _col.get("accent",   "A78BFA")   # Lavender-400
_HEX_VOID    = _col.get("void",     "09090B")   # near-black (cover band)
_HEX_MIST    = _col.get("mist",     "F5F3FF")   # Violet-50 (light surfaces)
_HEX_WARNING = _col.get("warning",  "FBBF24")
_HEX_SUCCESS = _col.get("success",  "4ADE80")

_HEX_Z50  = _zinc.get("50",  "FAFAFA")
_HEX_Z100 = _zinc.get("100", "F4F4F5")
_HEX_Z200 = _zinc.get("200", "E4E4E7")
_HEX_Z300 = _zinc.get("300", "D4D4D8")
_HEX_Z400 = _zinc.get("400", "A1A1AA")
_HEX_Z500 = _zinc.get("500", "71717A")
_HEX_Z600 = _zinc.get("600", "52525B")
_HEX_Z700 = _zinc.get("700", "3F3F46")
_HEX_Z900 = _zinc.get("900", "18181B")

# Callout component tokens
_HEX_CALL_INFO_BORDER = _call.get("info_border",     "A78BFA")
_HEX_CALL_INFO_BG     = _call.get("info_bg",         "F5F3FF")
_HEX_CALL_WARN_BORDER = _call.get("warning_border",  "FBBF24")
_HEX_CALL_WARN_BG     = _call.get("warning_bg",      "FFFBEB")
_HEX_CALL_OK_BORDER   = _call.get("success_border",  "4ADE80")
_HEX_CALL_OK_BG       = _call.get("success_bg",      "F0FDF4")

# Table component tokens
_HEX_TBL_HEADER_BG   = _tbl.get("header_bg",   "6D28D9")
_HEX_TBL_HEADER_TEXT = _tbl.get("header_text",  "FFFFFF")
_HEX_TBL_ALT_BG      = _tbl.get("row_alt_bg",   "F5F3FF")


# ─── python-docx ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ─── Color helper ─────────────────────────────────────────────────────────────
def _rgb(hex6: str) -> RGBColor:
    h = hex6.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


C_PRIMARY = _rgb(_HEX_PRIMARY)
C_ACCENT  = _rgb(_HEX_ACCENT)
C_VOID    = _rgb(_HEX_VOID)
C_MIST    = _rgb(_HEX_MIST)
C_Z400    = _rgb(_HEX_Z400)
C_Z500    = _rgb(_HEX_Z500)
C_Z600    = _rgb(_HEX_Z600)
C_Z700    = _rgb(_HEX_Z700)
C_Z900    = _rgb(_HEX_Z900)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Table / cell XML helpers ─────────────────────────────────────────────────
def _set_cell_shading(cell, hex6: str) -> None:
    h = hex6.lstrip("#")
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{h}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _remove_borders(table) -> None:
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}></w:tblPr>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_left_border(cell, hex6: str, sz: str = "24") -> None:
    h = hex6.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{h}"/>'
        f'</w:tcBorders>'
    ))


def _add_bottom_border(para, hex6: str, sz: str = "12") -> None:
    h = hex6.lstrip("#")
    para._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{h}"/>'
        f'</w:pBdr>'
    ))


# ─── Run helpers ──────────────────────────────────────────────────────────────
def body_run(para, text: str, *, bold: bool = False, italic: bool = False,
             color: RGBColor | None = None, size_pt: float = 9.5) -> None:
    """Add an Inter body run to an existing paragraph."""
    r = para.add_run(text)
    r.font.name   = FONT_BODY
    r.font.size   = Pt(size_pt)
    r.font.bold   = bold
    r.font.italic = italic
    r.font.color.rgb = color or C_Z600


def display_run(para, text: str, *, bold: bool = True, italic: bool = False,
                size_pt: float = 14, color: RGBColor | None = None) -> None:
    """Add a DM Serif Display run — for section titles and headings."""
    r = para.add_run(text)
    r.font.name   = FONT_DISPLAY
    r.font.size   = Pt(size_pt)
    r.font.bold   = bold
    r.font.italic = italic
    r.font.color.rgb = color or C_Z900


# ─── Document factory ─────────────────────────────────────────────────────────
def new_branded_doc() -> Document:
    """Return a Document pre-set with Strategitect page margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)
    return doc


# ─── Cover page ───────────────────────────────────────────────────────────────
def apply_cover(doc: Document, title: str, subtitle: str,
                date_str: str = "", client: str = "") -> None:
    """
    Obsidian editorial light cover.

    ┌─────────────────────────────────┐
    │  VIOLET BAND  ·  CONFIDENTIAL  │  ← Violet-700 band, white badge
    ├─────────────────────────────────┤
    │                                 │
    │  Title                          │  ← DM Serif Display 32pt Zinc-900
    │  ─────────────────────          │  ← Violet-700 rule
    │  Subtitle                       │  ← Inter 12pt Zinc-500
    │                                 │
    │  Prepared for │ Client          │  ← Metadata table
    │  Date         │ …               │
    │  Classification│ CONFIDENTIAL   │
    │  Prepared by  │ Strategitect    │
    │                                 │
    │  Architecture for PA Strategy   │  ← Slogan, italic, Zinc-400
    └─────────────────────────────────┘
    """
    # ── Header band ───────────────────────────────────────────────────────────
    band = doc.add_table(rows=1, cols=1)
    band.style = "Table Grid"
    _remove_borders(band)
    bc = band.cell(0, 0)
    _set_cell_shading(bc, _HEX_PRIMARY)

    tr = bc._tc.getparent()
    tr.get_or_add_trPr().append(
        parse_xml(f'<w:trHeight {nsdecls("w")} w:val="1440" w:hRule="exact"/>')
    )
    p_b = bc.paragraphs[0]
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_b.paragraph_format.space_before = Pt(10)
    r_b = p_b.add_run(COVER_BADGE)
    r_b.font.name = FONT_BODY
    r_b.font.size = Pt(8)
    r_b.font.bold = True
    r_b.font.color.rgb = C_WHITE

    # ── Spacers ───────────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    display_run(p_title, title, size_pt=32, color=C_Z900)

    # ── Violet rule ───────────────────────────────────────────────────────────
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after  = Pt(10)
    _add_bottom_border(p_rule, _HEX_PRIMARY)

    # ── Subtitle ──────────────────────────────────────────────────────────────
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(2)
    body_run(p_sub, subtitle, size_pt=12, color=C_Z500)

    # ── Spacers ───────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    # ── Metadata ──────────────────────────────────────────────────────────────
    rows: list[tuple[str, str]] = []
    if client:
        rows.append(("Prepared for", client))
    if date_str:
        rows.append(("Date", date_str))
    rows.append(("Classification", COVER_BADGE))
    rows.append(("Prepared by", COMPANY_NAME))
    meta_table(doc, rows)

    # ── Slogan ────────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    p_slogan = doc.add_paragraph()
    body_run(p_slogan, SLOGAN, italic=True, size_pt=8, color=C_Z400)

    doc.add_page_break()


# ─── Section headings ─────────────────────────────────────────────────────────
def section_heading(doc: Document, text: str, level: int = 1) -> None:
    """
    level=1 — DM Serif Display 14pt Zinc-900, violet ▌ prefix
    level=2 — DM Serif Display 11pt Zinc-900
    level=3 — Inter Bold 8pt Violet-700 UPPERCASE (eyebrow/label)
    """
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(4)
        r_bar = p.add_run("▌")
        r_bar.font.name  = FONT_BODY
        r_bar.font.size  = Pt(14)
        r_bar.font.color.rgb = C_PRIMARY
        display_run(p, f"  {text}", size_pt=14, bold=True, color=C_Z900)
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        display_run(p, text, size_pt=11, bold=True, color=C_Z900)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(1)
        body_run(p, text.upper(), bold=True, size_pt=8, color=C_PRIMARY)


# ─── Body paragraph ───────────────────────────────────────────────────────────
def body_para(doc: Document, text: str, *, bold: bool = False,
              italic: bool = False, indent_cm: float = 0,
              size_pt: float = 9.5, color: RGBColor | None = None,
              space_after_pt: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after_pt)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    body_run(p, text, bold=bold, italic=italic, color=color, size_pt=size_pt)


# ─── Bullet item ──────────────────────────────────────────────────────────────
def bullet_item(doc: Document, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.4)
    p.paragraph_format.space_after = Pt(3)
    prefix = "•" if level == 0 else "–"
    body_run(p, f"{prefix}  {text}", size_pt=9.5)


# ─── Callout box ─────────────────────────────────────────────────────────────
def callout(doc: Document, text: str, kind: str = "info") -> None:
    """
    Tinted callout with colored left border.
    kind: "info" | "warning" | "positive"
    """
    _map = {
        "info":     (_HEX_CALL_INFO_BORDER, _HEX_CALL_INFO_BG),
        "warning":  (_HEX_CALL_WARN_BORDER, _HEX_CALL_WARN_BG),
        "positive": (_HEX_CALL_OK_BORDER,   _HEX_CALL_OK_BG),
    }
    border_hex, bg_hex = _map.get(kind, _map["info"])

    t = doc.add_table(rows=1, cols=1)
    _remove_borders(t)
    cell = t.cell(0, 0)
    _set_cell_shading(cell, bg_hex)
    _set_left_border(cell, border_hex, sz="24")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    body_run(p, text, size_pt=9.5, color=C_Z700)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Metadata table ───────────────────────────────────────────────────────────
def meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Two-column label|value table — used for cover and document metadata."""
    t = doc.add_table(rows=len(rows), cols=2)
    _remove_borders(t)
    t.columns[0].width = Cm(4)
    t.columns[1].width = Cm(12)

    for i, (label, value) in enumerate(rows):
        lc, vc = t.cell(i, 0), t.cell(i, 1)
        _set_cell_shading(lc, _HEX_Z100)

        for cell, txt, is_label in ((lc, label, True), (vc, value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            body_run(p, txt,
                     bold=is_label, size_pt=8.5 if is_label else 9,
                     color=C_Z600 if is_label else C_Z700)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── Data table ───────────────────────────────────────────────────────────────
def data_table(doc: Document, headers: list[str],
               rows: list[list[str]]) -> None:
    """Branded data table — violet header row, mist alternating rows."""
    n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n)
    _remove_borders(t)

    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        _set_cell_shading(cell, _HEX_TBL_HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        body_run(p, h, bold=True, size_pt=8.5, color=C_WHITE)

    for i, row_data in enumerate(rows):
        bg = _HEX_TBL_ALT_BG if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = t.cell(i + 1, j)
            _set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            body_run(p, str(val), size_pt=9, color=C_Z600)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── Footer ───────────────────────────────────────────────────────────────────
def add_footer(doc: Document) -> None:
    """Add centered branded footer to all document sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_run(p, FOOTER_TEXT, size_pt=7.5, color=C_Z400)
