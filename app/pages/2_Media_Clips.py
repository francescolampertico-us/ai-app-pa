"""
Media Clips — Streamlit Page
==============================
Search Google News with Boolean queries, generate filtered clips report.
Includes embedded Clip Cleaner to fix paywalled articles in the generated report.
"""

import streamlit as st
import sys
import os
import json
import tempfile
import datetime
import subprocess
import re
from pathlib import Path

TOOLKIT_ROOT = Path(__file__).resolve().parent.parent.parent
TOOL_ROOT = TOOLKIT_ROOT / "tools" / "media_clips"
CLEANER_ROOT = TOOLKIT_ROOT / "tools" / "media_clip_cleaner"
sys.path.insert(0, str(TOOL_ROOT / "execution"))
sys.path.insert(0, str(CLEANER_ROOT / "execution"))

st.set_page_config(page_title="Media Clips", page_icon="📰", layout="wide")

st.title("📰 Media Clips Generator")
st.caption("v0.1.0  |  Risk: 🟡 Yellow  |  DiGiacomo: #1 Legislative Monitoring")
st.markdown(
    "Daily media monitoring tool. Searches Google News with Boolean queries, "
    "filters to trusted sources, deduplicates, and generates a formatted .docx report."
)

st.divider()


# ─── Helper: rebuild .docx from clips_data ───────────────────────────────────
def _build_report(clips_data, report_topic, report_date_str):
    """Build a .docx report and email body text from clips_data list."""
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from html import escape

    def _set_font(run, bold=False):
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "Calibri")

    def _add_hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        safe_text = escape(text)
        hyperlink = parse_xml(
            r'<w:hyperlink r:id="{}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            r' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            r'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0000FF"/>'
            r'<w:u w:val="single"/></w:rPr><w:t>{}</w:t></w:r></w:hyperlink>'.format(
                r_id, safe_text
            )
        )
        paragraph._p.append(hyperlink)

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    email_body_text = ""

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t1 = p_title.add_run(f"{report_topic}\n")
    _set_font(run_t1, bold=True)
    run_t2 = p_title.add_run(report_date_str)
    _set_font(run_t2, bold=True)
    doc.add_paragraph()

    # Index
    for idx, clip in enumerate(clips_data, 1):
        index_line = f"{idx}. {clip['source']}: {clip['title']} – {clip['date']}\n{clip.get('url', '')}\n"
        email_body_text += index_line + "\n"

        p = doc.add_paragraph()
        run_idx = p.add_run(f"{idx}. {clip['source']}: ")
        _set_font(run_idx, bold=True)
        if clip.get("url"):
            _add_hyperlink(p, clip["url"], clip["title"])
        else:
            run_title = p.add_run(clip["title"])
            _set_font(run_title)
        run_date = p.add_run(f" – {clip['date']}")
        _set_font(run_date)

    doc.add_page_break()
    email_body_text += "\n" + "=" * 20 + "\n\n"

    # Full clips
    for clip in clips_data:
        # Header block
        p_head = doc.add_paragraph()
        run_src = p_head.add_run(f"{clip['source']}\n")
        _set_font(run_src, bold=True)
        if clip.get("url"):
            _add_hyperlink(p_head, clip["url"], clip["title"])
        else:
            run_hl = p_head.add_run(clip["title"])
            _set_font(run_hl)
        p_head.add_run("\n")
        author = clip.get("author") or "Staff"
        run_auth = p_head.add_run(f"By {author}\n")
        _set_font(run_auth)
        run_dt = p_head.add_run(clip["date"])
        _set_font(run_dt)

        email_body_text += f"{clip['source']}\n{clip['title']}\n{clip.get('url', '')}\nBy {author}\n{clip['date']}\n\n"

        # Body text
        body_text = clip.get("extracted_text", "")
        if body_text:
            for para in body_text.split("\n"):
                para = para.strip()
                if not para:
                    continue
                p_body = doc.add_paragraph(para)
                p_body.paragraph_format.space_after = Pt(12)
                for run in p_body.runs:
                    _set_font(run)
                email_body_text += f"{para}\n\n"
        else:
            p_body = doc.add_paragraph("[PASTE FULL TEXT HERE]")
            _set_font(p_body.runs[0])

        email_body_text += "\n" + "-" * 10 + "\n\n"

    # Build HTML email body with clickable links
    html_body = '<div style="font-family:Calibri,sans-serif;font-size:14px;">\n'
    html_body += f"<p>Good morning,</p>\n"
    html_body += f"<p>Please find attached the {report_topic} for {report_date_str}.</p>\n"
    html_body += "<hr>\n"

    # Index with hyperlinked titles
    for idx, clip in enumerate(clips_data, 1):
        url = clip.get("url", "")
        if url:
            html_body += f'<p><b>{idx}. {clip["source"]}:</b> <a href="{url}">{clip["title"]}</a> &ndash; {clip["date"]}</p>\n'
        else:
            html_body += f'<p><b>{idx}. {clip["source"]}:</b> {clip["title"]} &ndash; {clip["date"]}</p>\n'

    html_body += "<hr>\n"

    # Full article bodies
    for clip in clips_data:
        url = clip.get("url", "")
        author = clip.get("author") or "Staff"
        html_body += f'<p><b>{clip["source"]}</b><br>\n'
        if url:
            html_body += f'<a href="{url}">{clip["title"]}</a><br>\n'
        else:
            html_body += f'{clip["title"]}<br>\n'
        html_body += f'By {author}<br>\n{clip["date"]}</p>\n'

        body_text = clip.get("extracted_text", "")
        if body_text:
            for para in body_text.split("\n"):
                para = para.strip()
                if para:
                    html_body += f"<p>{para}</p>\n"

        html_body += "<hr style='border:none;border-top:1px dashed #ccc;'>\n"

    html_body += "<p>Best regards</p>\n</div>"

    return doc, email_body_text, html_body


# ─── Inputs ──────────────────────────────────────────────────────────────────
topic = st.text_input("Report topic / title", placeholder="e.g., India Media Clips")
queries = st.text_area(
    "Search queries (one per line, Boolean syntax supported)",
    placeholder='"India" AND ("elections" OR "Modi")\n"New Delhi" AND "trade"',
    height=100,
)

col1, col2, col3 = st.columns(3)
with col1:
    period = st.selectbox("Search period", ["24h", "12h", "72h", "7d"], index=0)
with col2:
    target_date = st.date_input("Target date", value=datetime.date.today())
with col3:
    source_filter = st.selectbox(
        "Source filter",
        ["Mainstream media only", "All sources", "Custom"],
        index=0,
    )

custom_sources = None
if source_filter == "Custom":
    custom_sources = st.text_area(
        "Custom trusted domains (one per line)",
        placeholder="nytimes.com\npolitico.com\nyourcustomsource.com",
        height=80,
    )

since_date = st.text_input("Filter since (optional)", placeholder="YYYY-MM-DD HH:MM")


# ─── Step 1: Generate Clips Report ──────────────────────────────────────────
if topic and queries and st.button("Generate Clips Report", type="primary"):
    with st.spinner("Searching and generating clips..."):
        try:
            tmpdir = tempfile.mkdtemp()
            cmd = [
                sys.executable,
                str(TOOL_ROOT / "execution" / "generate_clips.py"),
                "--topic", topic,
                "--queries", ",".join(q.strip() for q in queries.strip().split("\n") if q.strip()),
                "--period", period,
                "--target-date", str(target_date),
                "--output-dir", tmpdir,
                "--no-email",
            ]
            if since_date:
                cmd.extend(["--since", since_date])
            if source_filter == "All sources":
                cmd.append("--all-sources")
            elif source_filter == "Custom" and custom_sources:
                domains = ",".join(d.strip() for d in custom_sources.strip().split("\n") if d.strip())
                cmd.extend(["--custom-sources", domains])

            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=120,
                cwd=str(TOOL_ROOT / "execution"),
            )

            if result.stdout:
                with st.expander("Pipeline log", expanded=False):
                    st.code(result.stdout)
            if result.stderr:
                with st.expander("Errors/warnings", expanded=False):
                    st.code(result.stderr)

            # Load article data JSON
            json_files = list(Path(tmpdir).rglob("*_data.json"))
            docx_files = list(Path(tmpdir).rglob("*.docx"))

            if json_files:
                with open(json_files[0], "r", encoding="utf-8") as jf:
                    clips_data = json.load(jf)
                st.session_state["clips_data"] = clips_data
                st.session_state["clips_topic"] = topic
                st.session_state["clips_date"] = str(target_date)
                st.session_state["clips_tmpdir"] = tmpdir

                n_missing = sum(1 for c in clips_data if not c.get("has_full_text"))
                st.success(
                    f"Found {len(clips_data)} articles. "
                    f"{n_missing} need full text (paywalled/missing)."
                )
            elif docx_files:
                st.success(f"Report generated: {docx_files[0].name}")
                with open(docx_files[0], "rb") as f:
                    st.download_button(
                        "Download .docx report", data=f.read(),
                        file_name=docx_files[0].name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            else:
                st.warning(
                    "No articles matched your queries. "
                    "Check the pipeline log for details."
                )

        except subprocess.TimeoutExpired:
            st.error("Search timed out after 2 minutes. Try narrower queries or a shorter period.")
        except Exception as e:
            st.error(f"Error: {e}")
            import traceback
            st.code(traceback.format_exc())

elif not topic or not queries:
    st.info("Enter a topic and search queries to generate a media clips report.")


# ─── Step 2: Review Articles & Fix Paywalled Clips ──────────────────────────
if st.session_state.get("clips_data"):
    clips_data = st.session_state["clips_data"]

    st.divider()
    st.subheader("📋 Review Articles")
    st.markdown(
        "Articles marked with ❌ are missing full text (paywalled). "
        "Open the link, copy the article, and use the Clip Cleaner below to fix them."
    )

    for i, clip in enumerate(clips_data):
        has_text = clip.get("has_full_text", False)
        icon = "✅" if has_text else "❌"
        label = f"{icon} {i+1}. {clip['source']}: {clip['title']}"
        with st.expander(label, expanded=False):
            if clip.get("url"):
                st.markdown(f"[Open article]({clip['url']})")
            st.caption(f"{clip.get('author', 'Staff')} · {clip['date']}")
            if has_text:
                preview = clip["extracted_text"][:500]
                if len(clip["extracted_text"]) > 500:
                    preview += "..."
                st.text(preview)
            else:
                st.warning("No full text — use Clip Cleaner below to add it.")

    # ─── Clip Cleaner ────────────────────────────────────────────────────────
    st.divider()
    st.subheader("🧹 Clip Cleaner")
    st.markdown(
        "Paste the raw article text for a paywalled article. "
        "After cleaning, select which article to update in the report."
    )

    # Build selector options: only articles missing text (+ option for all)
    article_options = {
        f"{i+1}. {c['source']}: {c['title']}": i
        for i, c in enumerate(clips_data)
    }

    target_article = st.selectbox(
        "Update which article?",
        options=list(article_options.keys()),
        key="target_article_select",
    )

    raw_paste = st.text_area(
        "Paste raw article text",
        placeholder="Copy the full article text from the webpage and paste it here...",
        height=200,
        key="clip_cleaner_input",
    )

    cleaner_mode = st.radio(
        "Cleaning mode",
        ["LLM (recommended)", "Local (rule-based)"],
        index=0,
        horizontal=True,
        key="cleaner_mode",
    )

    if raw_paste and st.button("Clean & Update Article", type="primary", key="clean_btn"):
        try:
            import importlib.util
            _spec = importlib.util.spec_from_file_location(
                "clean_clip",
                str(CLEANER_ROOT / "execution" / "clean_clip.py"),
            )
            _cc = importlib.util.module_from_spec(_spec)
            _spec.loader.exec_module(_cc)

            # Get the title of the selected article for better cleaning
            sel_idx = article_options[target_article]
            article_title = clips_data[sel_idx].get("title", "")

            if cleaner_mode.startswith("LLM"):
                with st.spinner("Cleaning with LLM..."):
                    try:
                        if not os.environ.get("OPENAI_API_KEY"):
                            st.error("OPENAI_API_KEY not set. Using local cleaner.")
                            raise RuntimeError("No API key")
                        cleaned = _cc.clean_clip_llm_openai(raw_paste, "gpt-4o-mini", title=article_title)
                        st.success("Cleaned with LLM.")
                    except Exception as llm_err:
                        st.warning(f"LLM failed ({llm_err}), falling back to local cleaner.")
                        cleaned = _cc.clean_clip(raw_paste)
            else:
                cleaned = _cc.clean_clip(raw_paste)

            ok, issues = _cc.validate_output(cleaned)
            if not ok:
                st.warning("Cleaned text has minor issues:")
                for issue in issues:
                    st.markdown(f"- {issue}")

            st.text_area("Cleaned output", value=cleaned, height=300, key="clip_cleaner_output")

            # Update the article in clips_data
            idx = article_options[target_article]
            st.session_state["clips_data"][idx]["extracted_text"] = cleaned
            st.session_state["clips_data"][idx]["has_full_text"] = True

            n_remaining = sum(1 for c in st.session_state["clips_data"] if not c.get("has_full_text"))
            st.success(
                f"Updated article {idx+1}. "
                f"{n_remaining} article(s) still missing full text."
            )
        except Exception as e:
            st.error(f"Cleaning error: {e}")
            import traceback
            st.code(traceback.format_exc())

    # ─── Step 3: Build Final Report ─────────────────────────────────────────
    st.divider()
    st.subheader("📄 Build Final Report")

    n_missing = sum(1 for c in clips_data if not c.get("has_full_text"))
    if n_missing > 0:
        st.info(f"{n_missing} article(s) still missing full text. You can still build the report.")

    if st.button("Build & Download Report", type="primary", key="build_final_report"):
        with st.spinner("Building report..."):
            try:
                report_topic = st.session_state.get("clips_topic", topic or "Media Clips")
                report_date_str = st.session_state.get("clips_date", str(datetime.date.today()))
                try:
                    from dateutil import parser as dp
                    date_formatted = dp.parse(report_date_str).strftime("%B %d, %Y")
                except Exception:
                    date_formatted = report_date_str

                doc, email_body, html_body = _build_report(clips_data, report_topic, date_formatted)

                # Filename matches email subject: "Topic - Month Day, Year.docx"
                subject = f"{report_topic} - {date_formatted}"
                fname = f"{subject}.docx"

                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name

                st.session_state["final_report_path"] = tmp_path
                st.session_state["final_report_fname"] = fname
                st.session_state["final_report_subject"] = subject
                st.session_state["final_html_body"] = html_body

                with open(tmp_path, "rb") as f:
                    st.download_button(
                        "Download Report (.docx)",
                        data=f.read(),
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_final_report",
                    )
                st.success(f"Report ready: {fname}")

                # Flag email body to rebuild with updated content
                st.session_state["_rebuild_email"] = True

            except Exception as e:
                st.error(f"Error building report: {e}")
                import traceback
                st.code(traceback.format_exc())

    # ─── Step 4: Prepare Email Draft ─────────────────────────────────────────
    if st.session_state.get("final_report_path"):
        st.divider()
        st.subheader("📧 Prepare Email Draft")

        # Build default email body text with clips and URLs
        if "email_body_draft" not in st.session_state or st.session_state.get("_rebuild_email"):
            subject = st.session_state.get("final_report_subject", "Media Clips")
            body_lines = [f"Good morning,\n", f"Please find attached the {subject}.\n"]
            body_lines.append("=" * 40 + "\n")
            for idx, clip in enumerate(clips_data, 1):
                url = clip.get("url", "")
                body_lines.append(f"{idx}. {clip['source']}: {clip['title']}")
                if url:
                    body_lines.append(f"   {url}")
                body_lines.append(f"   {clip.get('date', '')}")
                body_lines.append("")
            body_lines.append("=" * 40 + "\n")
            for clip in clips_data:
                url = clip.get("url", "")
                author = clip.get("author") or "Staff"
                body_lines.append(clip["source"])
                body_lines.append(clip["title"])
                if url:
                    body_lines.append(url)
                body_lines.append(f"By {author}")
                body_lines.append(clip["date"])
                body_lines.append("")
                body_text = clip.get("extracted_text", "")
                if body_text:
                    for para in body_text.split("\n"):
                        para = para.strip()
                        if para:
                            body_lines.append(para)
                            body_lines.append("")
                body_lines.append("-" * 20)
                body_lines.append("")
            body_lines.append("\nBest regards")
            st.session_state["email_body_draft"] = "\n".join(body_lines)
            st.session_state.pop("_rebuild_email", None)

        email_body_edited = st.text_area(
            "Email body (edit before sending)",
            value=st.session_state["email_body_draft"],
            height=400,
            key="email_body_editor",
        )
        # Save edits back
        st.session_state["email_body_draft"] = email_body_edited

        e_col1, e_col2 = st.columns(2)
        with e_col1:
            email_sender = st.text_input("Sender email", placeholder="you@domain.com", key="final_email_sender")
        with e_col2:
            email_recipient = st.text_input("Recipient email", placeholder="team@domain.com", key="final_email_recipient")

        if email_sender and email_recipient and st.button(
            "Prepare Email Draft (Mail.app)", type="primary", key="prepare_email_btn"
        ):
            try:
                tmp_path = st.session_state["final_report_path"]
                fname = st.session_state["final_report_fname"]
                subject = st.session_state["final_report_subject"]

                # Rename the temp file to match the subject-based filename
                import shutil
                named_path = os.path.join(os.path.dirname(tmp_path), fname)
                shutil.copy2(tmp_path, named_path)

                # Write body to temp file to avoid AppleScript escaping issues
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=".txt", mode="w", encoding="utf-8"
                ) as bf:
                    bf.write(email_body_edited)
                    body_path = bf.name

                subject_safe = subject.replace('"', '\\"')
                recipients = [r.strip() for r in email_recipient.split(",")]
                sender_prop = f'sender:"{email_sender}", '

                script = f'''
                set bodyFile to POSIX file "{body_path}"
                set bodyContent to read bodyFile as «class utf8»
                set recipientList to {{{", ".join(['"' + r + '"' for r in recipients])}}}
                set theSubject to "{subject_safe}"
                set theFile to POSIX file "{named_path}"

                tell application "Mail"
                    set newMessage to make new outgoing message with properties {{{sender_prop}subject:theSubject, content:bodyContent, visible:true}}
                    tell newMessage
                        make new attachment with properties {{file name:theFile}} at after the last paragraph
                        repeat with r in recipientList
                            make new to recipient at end of to recipients with properties {{address:r}}
                        end repeat
                    end tell
                    activate
                end tell
                '''
                subprocess.run(["osascript", "-e", script], check=True)
                st.success("Email draft prepared in Mail.app — review and send when ready.")
            except Exception as mail_err:
                st.error(f"Email draft failed: {mail_err}")
                import traceback
                st.code(traceback.format_exc())

    # Reset button
    if st.button("Start Over", key="reset_clips"):
        for key in ["clips_data", "clips_topic", "clips_date", "clips_tmpdir",
                     "final_report_path", "final_report_fname", "final_report_subject",
                     "final_html_body", "email_body_draft"]:
            st.session_state.pop(key, None)
        st.rerun()


# ─── Standalone Clip Cleaner ─────────────────────────────────────────────────
st.divider()
st.subheader("🧹 Standalone Clip Cleaner")
st.markdown(
    "Clean any article text independently — useful outside the clips report workflow."
)

standalone_paste = st.text_area(
    "Paste raw article text",
    placeholder="Copy the full article text from the webpage and paste it here...",
    height=200,
    key="standalone_cleaner_input",
)

sa_col1, sa_col2 = st.columns(2)
with sa_col1:
    standalone_mode = st.radio(
        "Cleaning mode",
        ["LLM (recommended)", "Local (rule-based)"],
        index=0,
        horizontal=True,
        key="standalone_cleaner_mode",
    )
with sa_col2:
    standalone_title = st.text_input(
        "Article headline (optional, helps remove title from output)",
        key="standalone_title",
    )

if standalone_paste and st.button("Clean Article", key="standalone_clean_btn"):
    try:
        import importlib.util
        _spec = importlib.util.spec_from_file_location(
            "clean_clip",
            str(CLEANER_ROOT / "execution" / "clean_clip.py"),
        )
        _cc = importlib.util.module_from_spec(_spec)
        _spec.loader.exec_module(_cc)

        if standalone_mode.startswith("LLM"):
            with st.spinner("Cleaning with LLM..."):
                try:
                    if not os.environ.get("OPENAI_API_KEY"):
                        st.error("OPENAI_API_KEY not set. Using local cleaner.")
                        raise RuntimeError("No API key")
                    cleaned = _cc.clean_clip_llm_openai(
                        standalone_paste, "gpt-4o-mini",
                        title=standalone_title or None,
                    )
                    st.success("Cleaned with LLM.")
                except Exception as llm_err:
                    st.warning(f"LLM failed ({llm_err}), falling back to local cleaner.")
                    cleaned = _cc.clean_clip(standalone_paste)
        else:
            cleaned = _cc.clean_clip(standalone_paste)

        ok, issues = _cc.validate_output(cleaned)
        if not ok:
            st.warning("Cleaned text has minor issues:")
            for issue in issues:
                st.markdown(f"- {issue}")

        st.text_area("Cleaned output", value=cleaned, height=300, key="standalone_output")
        st.download_button(
            "Download cleaned text",
            data=cleaned,
            file_name="cleaned_clip.txt",
            mime="text/plain",
            key="standalone_download",
        )
    except Exception as e:
        st.error(f"Cleaning error: {e}")
        import traceback
        st.code(traceback.format_exc())
